"""
Gera 3 documentos DOCX para homologação do Verus.AI V2:
1. Documentação do Sistema V2
2. Relatório de Atividades V2
3. Roadmap de Evolução

Uso:
    cd backend && source venv/bin/activate
    python ../scripts/generate_docs_v2.py
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs')


def set_style(doc):
    """Configura estilo base do documento."""
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for i in range(1, 5):
        h = doc.styles[f'Heading {i}']
        h.font.name = 'Calibri'
        h.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        if i == 1:
            h.font.size = Pt(18)
        elif i == 2:
            h.font.size = Pt(14)
        elif i == 3:
            h.font.size = Pt(12)


def add_cover(doc, title, subtitle, version, date):
    """Adiciona capa."""
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Versao {version} | {date}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Verus.AI / PRODERJ')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()


def add_table(doc, headers, rows):
    """Adiciona tabela formatada."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()


# ============================================================
# DOCUMENTO 1: Documentacao do Sistema V2
# ============================================================
def generate_system_doc():
    doc = Document()
    set_style(doc)
    add_cover(doc, 'Verus.AI', 'Documentacao do Sistema\nHomologacao V2', '2.0', '31/03/2026')

    # SUMARIO
    doc.add_heading('Sumario', level=1)
    sumario = [
        '1. Visao Geral',
        '2. Arquitetura do Sistema',
        '3. Stack Tecnologica',
        '4. Modulo 1 - Formulario Guiado',
        '5. Modulo 2 - Assistente Inteligente com IA',
        '6. Modulo 3 - Copilot (Chat IA)',
        '7. Importacao Cross-Blueprint (DOD -> ETP -> TR -> Minuta)',
        '8. Agent Tools (Ferramentas de Agentes)',
        '9. Pesquisa de Precos com IA',
        '10. Base de Conhecimento (Knowledge Base)',
        '11. Blueprints e Agentes IA',
        '12. Gestao de Documentos',
        '13. Gestao de Usuarios e Permissoes',
        '14. Configuracoes do Sistema',
        '15. Backend - Servicos Principais',
        '16. Backend - Fluxo de Geracao End-to-End',
        '17. Eventos SSE (Server-Sent Events)',
        '18. Painel Administrativo (Django Admin)',
        '19. Infraestrutura e Deploy',
        '20. Seguranca',
    ]
    for item in sumario:
        doc.add_paragraph(item, style='List Number')
    doc.add_page_break()

    # 1. VISAO GERAL
    doc.add_heading('1. Visao Geral', level=1)
    doc.add_paragraph(
        'O Verus.AI é uma plataforma de gestao documental inteligente desenvolvida para o PRODERJ '
        '(Centro de Tecnologia da Informacao e Comunicacao do Estado do Rio de Janeiro). '
        'O sistema automatiza a elaboracao de documentos de licitacao conforme a Lei 14.133/2021 '
        '(Nova Lei de Licitacoes) e normas estaduais do RJ.'
    )
    doc.add_paragraph(
        'A plataforma e composta por tres modulos principais:'
    )
    doc.add_paragraph('Modulo 1 - Formulario Guiado: Geracao de documentos padronizados por templates e formularios configuraveis, com assistencia opcional de IA campo a campo.', style='List Bullet')
    doc.add_paragraph('Modulo 2 - Assistente Inteligente: Geracao automatizada de documentos completos (ETP, DOD, TR, Editais, Minutas) utilizando IA com RAG (Retrieval-Augmented Generation) sobre bases de conhecimento de legislacao.', style='List Bullet')
    doc.add_paragraph('Modulo 3 - Copilot: Chat de IA conversacional para duvidas sobre licitacoes, legislacao e processos.', style='List Bullet')

    doc.add_heading('Documentos Suportados', level=2)
    add_table(doc,
        ['Tipo', 'Sigla', 'Descricao', 'Secoes'],
        [
            ['Documento de Oficializacao da Demanda', 'DOD', 'Formaliza a necessidade de contratacao', '3 secoes'],
            ['Estudo Tecnico Preliminar', 'ETP', 'Analise tecnica detalhada da contratacao', '29 secoes'],
            ['Termo de Referencia', 'TR', 'Especificacoes para o edital de licitacao', '17 secoes'],
            ['Minuta de Edital', 'Edital', 'Documento base do certame licitatorio', '10 secoes'],
            ['Minuta de Contrato', 'Contrato', 'Minutas de contratos administrativos', '6 secoes'],
        ]
    )

    doc.add_heading('Ambiente de Producao', level=2)
    add_table(doc,
        ['Item', 'Valor'],
        [
            ['URL', 'proderjdoc.bravonix.ia.br'],
            ['Backend', 'Django 5.2 + DRF + LangGraph'],
            ['Frontend', 'Next.js 14 + React 18 + TypeScript'],
            ['Banco de Dados', 'PostgreSQL 16 + pgvector'],
            ['Infra', 'Docker + Traefik + Cloudflare R2'],
            ['IA Principal', 'IBM watsonx (Mistral Large 2)'],
        ]
    )
    doc.add_page_break()

    # 2. ARQUITETURA
    doc.add_heading('2. Arquitetura do Sistema', level=1)
    doc.add_paragraph(
        'O sistema segue uma arquitetura de microservicos containerizados com Docker, '
        'orquestrados por Traefik como reverse proxy com SSL automatico.'
    )

    doc.add_heading('Diagrama de Componentes', level=2)
    doc.add_paragraph(
        'Frontend (Next.js) -> API REST (Django DRF) -> DynamicGraphBuilder (LangGraph) -> '
        'UnifiedLLMService (watsonx/Claude/GPT) + KnowledgeBaseService (RAG) -> '
        'PgVectorService -> PostgreSQL + pgvector'
    )

    doc.add_heading('Estrutura de Diretorios', level=2)
    doc.add_paragraph('backend/apps/intelligent_assistant/ - Sistema principal de geracao com IA', style='List Bullet')
    doc.add_paragraph('backend/apps/core/ - DocumentType, LLMProvider, AuditLog, AIAnalysisConfig', style='List Bullet')
    doc.add_paragraph('backend/apps/agents/ - EmbeddingService (watsonx E5-Large)', style='List Bullet')
    doc.add_paragraph('backend/apps/forms/ - FormTemplate, FormAssistant', style='List Bullet')
    doc.add_paragraph('backend/apps/documents/ - Historico de documentos', style='List Bullet')
    doc.add_paragraph('backend/apps/price_research/ - Pesquisa de precos (Serper + PNCP)', style='List Bullet')
    doc.add_paragraph('backend/apps/kb/ - Knowledge Base (legado)', style='List Bullet')
    doc.add_paragraph('frontend/src/app/(dashboard)/ - Paginas do dashboard', style='List Bullet')
    doc.add_paragraph('frontend/src/components/ - 90+ componentes reutilizaveis', style='List Bullet')
    doc.add_paragraph('frontend/src/hooks/ - 25+ hooks customizados', style='List Bullet')
    doc.add_page_break()

    # 3. STACK TECNOLOGICA
    doc.add_heading('3. Stack Tecnologica', level=1)
    add_table(doc,
        ['Camada', 'Tecnologias', 'Versao'],
        [
            ['Frontend', 'React, TypeScript, Next.js, TailwindCSS, shadcn/ui', '18 / 5.x / 14 / 3.x'],
            ['Backend', 'Django, Django REST Framework, LangGraph', '5.2 / 3.15 / 0.2'],
            ['Banco de Dados', 'PostgreSQL + extensao pgvector', '16 / 0.7'],
            ['IA / LLM', 'IBM watsonx (Mistral Large 2), Embeddings E5-Large (1024d)', '-'],
            ['Busca Web', 'Serper.dev (Google Search API)', '-'],
            ['Busca PNCP', 'API publica do Portal Nacional de Contratacoes Publicas', '-'],
            ['Storage', 'Cloudflare R2 (S3-compatible)', '-'],
            ['Filas', 'Redis + Celery', '7.x / 5.x'],
            ['Infra', 'Docker, Traefik, Gunicorn', '-'],
            ['Markdown', 'ReactMarkdown + rehypeRaw + remarkGfm', '-'],
        ]
    )
    doc.add_page_break()

    # 4. MODULO 1 - FORMULARIO GUIADO
    doc.add_heading('4. Modulo 1 - Formulario Guiado', level=1)
    doc.add_paragraph('Rota: /dashboard/forms')
    doc.add_paragraph(
        'O modulo de Formulario Guiado permite gerar documentos padronizados a partir de '
        'templates configurados com campos especificos. O servidor preenche um formulario '
        'estruturado e o sistema gera o documento formatado automaticamente.'
    )

    doc.add_heading('Funcionalidades', level=2)
    add_table(doc,
        ['Funcionalidade', 'Descricao'],
        [
            ['Templates configuraveis', 'Cada tipo de documento tem seu template com campos especificos'],
            ['Campos condicionais', 'Campos que aparecem com base em respostas anteriores'],
            ['Validacao em tempo real', 'Campos obrigatorios validados antes de permitir geracao'],
            ['Assistencia IA por campo', 'Agentes de IA podem sugerir conteudo para campos individuais'],
            ['Exportacao PDF/DOCX', 'Documento gerado em formato profissional com layout configuravel'],
            ['Historico de versoes', 'Cada geracao e salva e pode ser consultada posteriormente'],
        ]
    )
    doc.add_page_break()

    # 5. MODULO 2 - ASSISTENTE INTELIGENTE
    doc.add_heading('5. Modulo 2 - Assistente Inteligente com IA', level=1)
    doc.add_paragraph('Rota: /dashboard/intelligent-assistant')
    doc.add_paragraph(
        'O Assistente Inteligente e o modulo central do Verus.AI. Utiliza agentes de IA '
        'especializados para gerar documentos completos de licitacao com base em Blueprints '
        'configuraveis, bases de conhecimento legislativa e documentos de referencia do usuario.'
    )

    doc.add_heading('Fluxo de Geracao (6 Fases)', level=2)

    doc.add_heading('Fase 1 - Selecao do Blueprint e Objetivo', level=3)
    doc.add_paragraph('O usuario seleciona o tipo de documento (ETP, DOD, TR, etc.) e informa o objetivo da contratacao. Opcionalmente, pode selecionar um DOD existente para importar conteudo (fluxo DOD -> ETP).')

    doc.add_heading('Fase 2 - Upload de Documentos de Referencia', level=3)
    doc.add_paragraph('Upload de PDFs, DOCX ou ODT como documentos de referencia. O sistema extrai texto, divide em chunks de 1000 caracteres com 200 de overlap, gera embeddings vetoriais (1024 dimensoes) e indexa no PostgreSQL via pgvector. Esses documentos ficam disponiveis para consulta RAG durante a geracao.')

    doc.add_heading('Fase 3 - Geracao do Documento', level=3)
    doc.add_paragraph('O usuario seleciona quais secoes gerar. Secoes com campos estruturados (ex: Responsavel, Equipe) sao preenchidas manualmente. Secoes com sub-secoes condicionais permitem escolher entre "detalhar" ou "nao se aplica". A geracao e feita via streaming (SSE) com progresso em tempo real.')
    doc.add_paragraph('Para cada secao, o pipeline executa: Geracao (agente gerador) -> Validacao (agente validador com score 0-100) -> Refinamento (se score < 70, ate 3 tentativas).')

    doc.add_heading('Fase 4 - Avaliacao das Secoes', level=3)
    doc.add_paragraph('Apos a geracao, cada secao exibe o texto gerado, score de confianca, erros e avisos. O usuario pode aprovar, editar manualmente, ou regenerar com feedback especifico.')

    doc.add_heading('Fase 5 - Resultado Final', level=3)
    doc.add_paragraph('Preview completo do documento com todas as secoes. Exportacao disponivel em PDF (formatado com logo, cores e layout do Blueprint), DOCX (editavel no Word) e ODT (editavel no LibreOffice).')

    doc.add_heading('Fase 6 - Historico', level=3)
    doc.add_paragraph('Acesso a sessoes anteriores pelo painel lateral. Visualizacao e restauracao de documentos gerados anteriormente.')

    doc.add_heading('Funcionalidades do Modulo', level=2)
    add_table(doc,
        ['Funcionalidade', 'Descricao'],
        [
            ['Upload de referencias', 'PDF, DOCX, ODT processados e indexados automaticamente'],
            ['Geracao por secao', 'Cada secao gerada por agente IA especializado'],
            ['RAG hierarquico', 'Busca em 4 camadas: Session + Global + Blueprint + Agente'],
            ['Validacao automatica', 'Agente validador avalia qualidade e conformidade de cada secao'],
            ['Regeneracao seletiva', 'Regenere apenas a secao que precisa de ajuste, com feedback'],
            ['Auto-aprendizagem', 'Feedbacks alimentam KB Camada 3 para melhorar futuras geracoes'],
            ['Streaming em tempo real', 'Progresso via SSE (Server-Sent Events)'],
            ['Sub-secoes condicionais', 'Logica OU: "detalhar" ou "nao se aplica" por sub-item'],
            ['Campos estruturados', 'Formularios por secao (text, date, select, array, etc.)'],
            ['Importacao cross-blueprint', 'DOD -> ETP -> TR -> Minuta (conteudo de documentos anteriores)'],
            ['Agent Tools', 'Ferramentas externas (Serper, PNCP) executadas antes da geracao'],
            ['Scores arredondados', 'Todos os scores com max 2 casas decimais (sem dizimas)'],
        ]
    )
    doc.add_page_break()

    # 6. MODULO 3 - COPILOT
    doc.add_heading('6. Modulo 3 - Copilot (Chat IA)', level=1)
    doc.add_paragraph('Rota: /dashboard/copilot')
    doc.add_paragraph(
        'Chat conversacional com IA para duvidas sobre licitacoes, legislacao, '
        'processos e uso do sistema. Utiliza o mesmo UnifiedLLMService com acesso '
        'as bases de conhecimento configuradas.'
    )
    doc.add_heading('Funcionalidades', level=2)
    doc.add_paragraph('Chat em tempo real com IA (streaming)', style='List Bullet')
    doc.add_paragraph('Historico de conversas por sessao', style='List Bullet')
    doc.add_paragraph('Edicao de mensagens enviadas (textarea auto-expansivel)', style='List Bullet')
    doc.add_paragraph('Renderizacao de Markdown (tabelas, listas, codigo)', style='List Bullet')
    doc.add_paragraph('Largura responsiva das mensagens (max 90%)', style='List Bullet')
    doc.add_page_break()

    # 7. IMPORTACAO CROSS-BLUEPRINT
    doc.add_heading('7. Importacao Cross-Blueprint', level=1)
    doc.add_paragraph(
        'O sistema permite importar conteudo de documentos anteriores para pre-popular secoes '
        'de novos documentos, seguindo o fluxo real de licitacao: DOD -> ETP -> TR -> Minuta.'
    )

    doc.add_heading('Fluxo de Importacao', level=2)
    add_table(doc,
        ['Origem', 'Destino', 'Mapeamento', 'Tipo'],
        [
            ['DOD Secao 3 (Demanda)', 'ETP Secao 1 (Descricao da Necessidade)', 'SectionImportConfig', 'mixed'],
            ['ETP (varias secoes)', 'TR (secoes correspondentes)', 'SectionImportConfig', 'mixed'],
            ['TR (varias secoes)', 'Minuta de Edital', 'SectionImportConfig', 'mixed'],
            ['TR (varias secoes)', 'Minuta de Contrato', 'SectionImportConfig', 'mixed'],
        ]
    )

    doc.add_heading('Como Funciona', level=2)
    doc.add_paragraph('1. O usuario seleciona o blueprint destino (ex: ETP)', style='List Number')
    doc.add_paragraph('2. O sistema lista sessoes do documento fonte disponiveis (ex: DODs concluidos)', style='List Number')
    doc.add_paragraph('3. O usuario seleciona uma sessao fonte', style='List Number')
    doc.add_paragraph('4. O sistema busca conteudo das secoes mapeadas via SectionImportConfig', style='List Number')
    doc.add_paragraph('5. O conteudo importado pre-preenche o objetivo e as secoes correspondentes', style='List Number')
    doc.add_paragraph('6. O agente gerador recebe o conteudo importado como contexto prioritario', style='List Number')

    doc.add_heading('Servicos de Importacao', level=2)
    add_table(doc,
        ['Servico', 'Arquivo', 'Funcao'],
        [
            ['ETPImportService', 'services/etp_import_service.py', 'Lista DODs e importa DOD -> ETP'],
            ['TRImportService', 'services/tr_import_service.py', 'Lista ETPs e importa ETP -> TR'],
            ['MinutaImportService', 'services/minuta_import_service.py', 'Lista TRs e importa TR -> Minutas'],
        ]
    )
    doc.add_page_break()

    # 8. AGENT TOOLS
    doc.add_heading('8. Agent Tools (Ferramentas de Agentes)', level=1)
    doc.add_paragraph(
        'Agentes de IA podem ter ferramentas externas vinculadas que sao executadas ANTES '
        'da chamada ao LLM. Os resultados das ferramentas sao injetados no prompt como contexto '
        'adicional, permitindo que o agente use dados reais de mercado.'
    )

    doc.add_heading('Ferramentas Disponiveis', level=2)
    add_table(doc,
        ['Ferramenta', 'Tipo', 'Descricao', 'Uso'],
        [
            ['Serper.dev (Google Search)', 'web_search', 'Busca no Google via API Serper.dev', 'Pesquisa de precos, fornecedores, mercado'],
            ['PNCP Search', 'pncp_search', 'Busca no Portal Nacional de Contratacoes Publicas', 'Licitacoes anteriores, precos de referencia'],
        ]
    )

    doc.add_heading('Arquitetura', level=2)
    doc.add_paragraph('1. AgentTool (model): define a ferramenta (nome, tipo, config default)', style='List Number')
    doc.add_paragraph('2. AgentToolLink (model): vincula ferramenta ao agente com prioridade e config customizada', style='List Number')
    doc.add_paragraph('3. AgentToolsService: orquestra execucao das ferramentas antes da chamada LLM', style='List Number')
    doc.add_paragraph('4. Mini-DAG: secao 5 executa tools e salva resultados; secao 6 herda via tool_results', style='List Number')

    doc.add_heading('Fluxo de Execucao', level=2)
    doc.add_paragraph(
        'Antes de chamar o LLM, o DynamicGraphBuilder verifica se o agente tem tools vinculadas. '
        'Se sim, o AgentToolsService usa o LLM para gerar queries otimizadas (separadas para Google '
        'e PNCP), executa as buscas em paralelo, formata os resultados como Markdown e injeta '
        'no prompt do agente como secao "DADOS DE MERCADO".'
    )
    doc.add_page_break()

    # 9. PESQUISA DE PRECOS
    doc.add_heading('9. Pesquisa de Precos com IA', level=1)
    doc.add_paragraph('Rota: /dashboard/price-research')
    doc.add_paragraph(
        'Modulo de pesquisa de precos que combina busca na web (Serper.dev/Google) '
        'com busca no PNCP (Portal Nacional de Contratacoes Publicas) para obter '
        'precos de referencia de mercado.'
    )

    doc.add_heading('Funcionalidades', level=2)
    doc.add_paragraph('Busca web via Serper.dev (Google Search API) - custo US$ 0.001/busca', style='List Bullet')
    doc.add_paragraph('Busca PNCP com enriquecimento de itens (detalhes de preco por item)', style='List Bullet')
    doc.add_paragraph('Analise automatica por IA (configuravel via AIAnalysisConfig)', style='List Bullet')
    doc.add_paragraph('Modal de analise com Markdown renderizado (tabelas, listas)', style='List Bullet')
    doc.add_paragraph('Exportacao de resultados', style='List Bullet')

    doc.add_heading('AIAnalysisConfig', level=2)
    doc.add_paragraph(
        'A analise de editais e configuravel via modelo AIAnalysisConfig no Django Admin. '
        'Permite definir: provider/modelo LLM, system_prompt personalizado, perguntas de analise '
        '(JSONField), temperatura e max_tokens.'
    )
    doc.add_page_break()

    # 10. BASE DE CONHECIMENTO
    doc.add_heading('10. Base de Conhecimento (Knowledge Base)', level=1)
    doc.add_paragraph('Rota: /dashboard/knowledge-base')

    doc.add_heading('Arquitetura de 4 Camadas', level=2)
    add_table(doc,
        ['Camada', 'Nome', 'Escopo', 'Exemplo'],
        [
            ['0', 'Session Docs', 'Documentos da sessao do usuario', 'PDFs enviados como referencia'],
            ['1', 'Global', 'Todos os agentes e blueprints', 'Lei 14.133, LGPD, Decreto 48.997'],
            ['2', 'Blueprint', 'Especifico de um tipo de documento', 'KB do ETP RJ, KB do DOD RJ'],
            ['3', 'Agente', 'Melhores resultados de um agente', 'Auto-aprendizagem via feedbacks'],
        ]
    )

    doc.add_heading('AgentKnowledgeBaseLink (Vinculo Avancado)', level=2)
    doc.add_paragraph(
        'O sistema utiliza um modelo intermediario (AgentKnowledgeBaseLink) que permite '
        'configuracao granular de como cada agente consulta cada KB:'
    )
    add_table(doc,
        ['Campo', 'Tipo', 'Descricao'],
        [
            ['priority', 'IntegerField', 'Ordem de consulta (0 = primeiro)'],
            ['purpose', 'CharField', 'Proposito: examples, evaluation, normative, context, reference'],
            ['top_k', 'IntegerField', 'Chunks retornados por KB (default: 5)'],
            ['min_similarity', 'FloatField', 'Threshold de relevancia (default: 0.6)'],
            ['include_summary', 'BooleanField', 'Incluir resumo interpretativo com chunks'],
            ['selected_sources', 'JSONField', 'Filtrar por fontes especificas'],
            ['instruction', 'TextField', 'Instrucao customizada para o agente sobre esta KB'],
        ]
    )

    doc.add_heading('Metodo Principal: query_by_links()', level=2)
    doc.add_paragraph(
        'O KnowledgeBaseService.query_by_links() consulta KBs na ordem de prioridade definida '
        'pelos links, aplicando configuracoes individuais (top_k, min_similarity) por KB. '
        'Inclui fallback automatico para camadas globais e de blueprint se nao houver link explicito.'
    )

    doc.add_heading('Bases em Producao', level=2)
    add_table(doc,
        ['Base de Conhecimento', 'Camada', 'Documentos', 'Embeddings'],
        [
            ['KB Global', 'Global', 'Lei 14.133, LGPD, Decreto 48.997, Lei 4.480, PEDTIC', '597'],
            ['KB - ETP RJ - PRODERJ', 'Blueprint', 'Sessao 1 - ETP.pdf', '129'],
        ]
    )
    doc.add_page_break()

    # 11. BLUEPRINTS E AGENTES IA
    doc.add_heading('11. Blueprints e Agentes IA', level=1)
    doc.add_paragraph('Rota: /dashboard/blueprints')

    doc.add_heading('O que e um Blueprint', level=2)
    doc.add_paragraph(
        'Um Blueprint define a estrutura completa de um tipo de documento: quais secoes o compoem, '
        'qual agente IA gera cada secao, qual agente valida, quais bases de conhecimento sao '
        'consultadas, e como o PDF final e formatado (cores, fontes, margens, capa).'
    )

    doc.add_heading('Blueprints em Producao', level=2)
    add_table(doc,
        ['Blueprint', 'Tipo', 'Secoes', 'Geradores', 'Validadores'],
        [
            ['ETP RJ - PRODERJ', 'ETP', '29', '29', '29'],
            ['DOD RJ - PRODERJ', 'DOD', '3', '3', '3'],
            ['TR RJ - PRODERJ', 'TR', '17', '17', '17'],
            ['Edital - PRODERJ', 'Edital', '10', '10', '10'],
            ['Minuta de Contrato - PRODERJ', 'Contrato', '6', '6', '6'],
        ]
    )

    doc.add_heading('Tipos de Agente', level=2)
    add_table(doc,
        ['Tipo', 'Funcao', 'Quando usar'],
        [
            ['Generator', 'Gera conteudo da secao', 'Toda secao que precisa de IA'],
            ['Validator', 'Avalia qualidade (score 0-100)', 'Secoes que exigem conformidade legal'],
            ['Refiner', 'Melhora conteudo com feedback', 'Automatico quando validacao falha'],
        ]
    )

    doc.add_heading('Variaveis Disponiveis nos Prompts', level=2)
    add_table(doc,
        ['Variavel', 'Conteudo'],
        [
            ['{{objective}}', 'Objetivo informado pelo usuario'],
            ['{{objective_summary}}', 'Resumo do objetivo (para queries RAG longas, max 512 tokens)'],
            ['{{section_name}}', 'Nome da secao sendo gerada'],
            ['{{section_number}}', 'Numero da secao'],
            ['{{context}}', 'Objetivo + instrucoes do Blueprint'],
            ['{{previous_sections}}', 'Conteudo das secoes ja geradas (dependencias)'],
            ['{{current_content}}', 'Conteudo atual (para refinamento)'],
            ['{{improvements}}', 'Feedback do validador'],
            ['{{input}} / {{user_input}}', 'Dados de campos estruturados de sub-secoes'],
            ['{{fields_data}}', 'Dados de campos estruturados formatados'],
            ['{{etp_imported_content}}', 'Conteudo importado de outro documento'],
            ['{{tool_results}}', 'Resultados de Agent Tools (pesquisa de mercado, PNCP)'],
        ]
    )
    doc.add_page_break()

    # 12. GESTAO DE DOCUMENTOS
    doc.add_heading('12. Gestao de Documentos', level=1)
    doc.add_paragraph('Rota: /dashboard/documents')
    doc.add_paragraph(
        'Listagem unificada de todos os documentos gerados pelo usuario, com filtros por status, '
        'origem e tipo. Permite visualizar, baixar (PDF/DOCX/ODT), editar e excluir documentos.'
    )
    doc.add_paragraph(
        'Os documentos finalizados sao armazenados no Cloudflare R2 (S3-compatible) via tasks '
        'Celery assincronas com retry automatico (3 tentativas).'
    )
    doc.add_page_break()

    # 13. GESTAO DE USUARIOS
    doc.add_heading('13. Gestao de Usuarios e Permissoes', level=1)
    doc.add_paragraph('Rota: /dashboard/users')

    doc.add_heading('Perfis de Acesso', level=2)
    add_table(doc,
        ['Role', 'Descricao', 'Permissoes Principais'],
        [
            ['Superadmin', 'Acesso total ao sistema', 'Tudo'],
            ['Manager', 'Gerente de equipe', 'Gerenciar usuarios, documentos, configuracoes'],
            ['Analyst', 'Analista / Elaborador', 'Usar assistente IA, formularios, geradores'],
            ['Reviewer', 'Revisor', 'Revisar e aprovar documentos'],
            ['Viewer', 'Visualizador', 'Acesso somente leitura'],
        ]
    )
    doc.add_page_break()

    # 14. CONFIGURACOES
    doc.add_heading('14. Configuracoes do Sistema', level=1)

    doc.add_heading('14.1. Perfil do Usuario', level=2)
    doc.add_paragraph('Rota: /dashboard/settings/profile - Alterar dados pessoais, senha e preferencia de tema (claro/escuro).')

    doc.add_heading('14.2. Identidade Visual', level=2)
    doc.add_paragraph('Rota: /dashboard/settings/brand - Nome, tagline, logos (modo claro/escuro), favicon, cores primaria/secundaria/destaque. Preview em tempo real.')

    doc.add_heading('14.3. Perfis e Permissoes', level=2)
    doc.add_paragraph('Rota: /dashboard/settings/roles - Configurar permissoes por role, criar/editar perfis customizados.')
    doc.add_page_break()

    # 15. BACKEND - SERVICOS
    doc.add_heading('15. Backend - Servicos Principais', level=1)

    services = [
        ('DynamicGraphBuilder', 'services/dynamic_graph_builder.py',
         'Orquestrador principal. Constroi pipelines de geracao com LangGraph. Gerencia estado, '
         'executa Agent Tools, injeta contexto RAG, gerencia tentativas de refinamento. '
         'Suporta secoes com campos estruturados, sub-secoes condicionais, importacao cross-blueprint.'),
        ('KnowledgeBaseService', 'services/knowledge_base_service.py',
         'Busca semantica nas 4 camadas de KB. Metodo principal: query_by_links() com prioridade '
         'e configuracao por KB via AgentKnowledgeBaseLink. Fallback automatico para camadas nao cobertas.'),
        ('PgVectorService', 'services/pgvector_service.py',
         'Gerencia embeddings e busca vetorial. Gera embeddings via IBM watsonx E5-Large (1024d). '
         'Busca por similaridade de cosseno. Suporta batch processing.'),
        ('UnifiedLLMService', 'services/llm_provider_service.py',
         'Interface unificada para LLMs multi-provider. Suporta watsonx (Mistral Large), '
         'Anthropic (Claude), OpenAI (GPT). Credenciais gerenciadas via Django Admin.'),
        ('AgentToolsService', 'services/agent_tools_service.py',
         'Orquestra execucao de ferramentas externas antes da chamada LLM. Usa LLM para gerar '
         'queries otimizadas. Executa tools em ordem de prioridade.'),
        ('ETPImportService', 'services/etp_import_service.py',
         'Importa conteudo DOD -> ETP. Lista sessoes DOD disponiveis, mapeia secoes via SectionImportConfig.'),
        ('TRImportService', 'services/tr_import_service.py',
         'Importa conteudo ETP -> TR. Mesmo padrao do ETPImportService.'),
        ('MinutaImportService', 'services/minuta_import_service.py',
         'Importa conteudo TR -> Minutas de Edital e Contrato.'),
        ('DocumentProcessorService', 'services/document_processor.py',
         'Extrai texto de PDF (pypdf), DOCX (python-docx), ODT (odfpy), TXT. Validacao de magic bytes.'),
        ('DocumentExportService', 'services/document_export_service.py',
         'Gera PDF/DOCX/ODT formatados com configuracoes do Blueprint (cores, fontes, capa, cabecalho).'),
        ('WebSearchService', 'price_research/services/web_search_service.py',
         'Busca web via Serper.dev (Google Search API). Busca PNCP com enriquecimento de itens.'),
    ]

    for name, path, desc in services:
        doc.add_heading(name, level=2)
        doc.add_paragraph(f'Arquivo: {path}')
        doc.add_paragraph(desc)

    doc.add_page_break()

    # 16. FLUXO DE GERACAO
    doc.add_heading('16. Backend - Fluxo de Geracao End-to-End', level=1)
    doc.add_paragraph(
        'O fluxo completo de geracao de um documento segue os seguintes passos:'
    )
    steps = [
        'Usuario clica "Gerar" -> POST /generate-etp-dynamic-stream/',
        'Criar GenerationSession (status=generating) e SectionGeneration para cada secao',
        'Inicializar DynamicGraphBuilder com Blueprint',
        'Para cada secao: verificar se tem campos estruturados (formatar markdown, score=100, pular LLM)',
        'Se secao com IA: carregar agente gerador, verificar Agent Tools vinculadas',
        'Se tem Agent Tools: AgentToolsService gera queries, executa buscas, injeta resultados no prompt',
        'Consultar RAG (query_by_links) se use_rag=true',
        'Montar prompt: system_prompt + user_prompt com {{variaveis}} + contexto RAG + tool_results',
        'Chamar LLM (watsonx/Claude/GPT) via UnifiedLLMService',
        'Validar com agente validador (score 0-100, arredondado 2 casas)',
        'Se score < 70 e tentativas < max: refinar com feedback do validador e repetir',
        'Se valido: salvar SectionGeneration, emitir SSE section_validated',
        'Ao finalizar todas as secoes: calcular average_score, emitir SSE completed',
        'Persistir documento final (GeneratedDocument) e upload para R2 via Celery task',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_page_break()

    # 17. EVENTOS SSE
    doc.add_heading('17. Eventos SSE (Server-Sent Events)', level=1)
    doc.add_paragraph(
        'A geracao de documentos utiliza streaming via SSE para feedback em tempo real ao usuario.'
    )
    add_table(doc,
        ['Evento', 'Quando', 'Dados'],
        [
            ['started', 'Inicio da geracao', 'session_id, blueprint'],
            ['section_start', 'Nova secao iniciando', 'section_number, section_name'],
            ['generating', 'Conteudo sendo gerado', 'section_number, progress'],
            ['section_content', 'Conteudo pronto', 'section_number, content'],
            ['section_validated', 'Validacao concluida', 'section_number, score, is_valid'],
            ['llm_usage', 'Auditoria de tokens', 'provider, model, input_tokens, output_tokens, duration_ms'],
            ['progress', 'Atualizacao de progresso', 'completed, total, percentage'],
            ['completed', 'Geracao finalizada', 'session_id, total_tokens, cost, average_score'],
            ['error', 'Erro geral', 'message'],
        ]
    )
    doc.add_page_break()

    # 18. ADMIN
    doc.add_heading('18. Painel Administrativo (Django Admin)', level=1)
    doc.add_paragraph('URL: /admin/')

    doc.add_heading('Modelos Gerenciaveis', level=2)
    add_table(doc,
        ['Modelo', 'App', 'Funcao'],
        [
            ['DocumentBlueprint', 'Intelligent Assistant', 'Estrutura do documento (secoes, agentes, layout)'],
            ['BlueprintSection', 'Intelligent Assistant', 'Secoes do blueprint com agentes e dependencias'],
            ['SectionAgentConfig', 'Intelligent Assistant', 'Config de agentes IA (prompts, modelo, RAG, tools)'],
            ['AgentTool', 'Intelligent Assistant', 'Ferramentas externas (Serper, PNCP)'],
            ['AgentToolLink', 'Intelligent Assistant', 'Vinculo agente <-> ferramenta com prioridade'],
            ['KnowledgeBase', 'Intelligent Assistant', 'Bases de conhecimento (4 camadas)'],
            ['AgentKnowledgeBaseLink', 'Intelligent Assistant', 'Vinculo agente <-> KB com prioridade e config'],
            ['SectionImportConfig', 'Intelligent Assistant', 'Mapeamento de importacao cross-blueprint'],
            ['GenerationSession', 'Intelligent Assistant', 'Sessoes de geracao de documentos'],
            ['LLMAuditLog', 'Intelligent Assistant', 'Auditoria de chamadas LLM (tokens, custo)'],
            ['LLMProvider', 'Core', 'Credenciais dos providers de IA'],
            ['LLMModel', 'Core', 'Modelos disponiveis por provider'],
            ['DocumentType', 'Core', 'Catalogo de tipos de documento'],
            ['AIAnalysisConfig', 'Core', 'Config de analise IA para editais'],
        ]
    )
    doc.add_page_break()

    # 19. INFRA
    doc.add_heading('19. Infraestrutura e Deploy', level=1)

    doc.add_heading('Containers Docker', level=2)
    add_table(doc,
        ['Container', 'Imagem', 'Funcao'],
        [
            ['backend', 'mqmaellson39/verus-backend:proderjdoc', 'Django + Gunicorn (API)'],
            ['frontend', 'mqmaellson39/verus-ai-frontend:proderjdoc', 'Next.js (SSR)'],
            ['postgres', 'pgvector/pgvector:pg16', 'PostgreSQL 16 + pgvector'],
            ['redis', 'redis:7-alpine', 'Broker para Celery'],
            ['celery-worker', 'verus-backend', 'Worker Celery (tasks assincronas)'],
            ['traefik', 'traefik:v2.11', 'Reverse proxy + SSL automatico'],
        ]
    )

    doc.add_heading('Build e Deploy', level=2)
    doc.add_paragraph('Backend:')
    doc.add_paragraph('docker buildx build --platform linux/amd64 -t mqmaellson39/verus-backend:proderjdoc -f Dockerfile.prod --push .', style='List Bullet')
    doc.add_paragraph('Frontend:')
    doc.add_paragraph('docker buildx build --platform linux/amd64 -t mqmaellson39/verus-ai-frontend:proderjdoc -f Dockerfile --push .', style='List Bullet')

    doc.add_heading('Storage (Cloudflare R2)', level=2)
    doc.add_paragraph(
        'Documentos finalizados (PDF/DOCX) sao salvos no Cloudflare R2 via tasks Celery '
        'assincronas com retry automatico (3 tentativas, delay de 10s). Configurado via '
        'STORAGES dict no Django 5.2+.'
    )
    doc.add_page_break()

    # 20. SEGURANCA
    doc.add_heading('20. Seguranca', level=1)
    doc.add_paragraph('Autenticacao JWT (JSON Web Token) via Django REST Framework', style='List Bullet')
    doc.add_paragraph('Permissoes granulares por funcionalidade e role', style='List Bullet')
    doc.add_paragraph('CORS configurado para dominios permitidos', style='List Bullet')
    doc.add_paragraph('Validacao de magic bytes em uploads (prevencao de uploads maliciosos)', style='List Bullet')
    doc.add_paragraph('SSL automatico via Traefik + Lets Encrypt', style='List Bullet')
    doc.add_paragraph('Credenciais LLM gerenciadas no banco (nao em .env)', style='List Bullet')
    doc.add_paragraph('Rate limiting em endpoints criticos', style='List Bullet')
    doc.add_paragraph('Auditoria completa de chamadas LLM (tokens, custo, duracao)', style='List Bullet')

    # SALVAR
    path = os.path.join(OUTPUT_DIR, 'verus_ai_Documentacao_Sistema_V2.docx')
    doc.save(path)
    print(f'[OK] {path}')
    return path


# ============================================================
# DOCUMENTO 2: Relatorio de Atividades V2
# ============================================================
def generate_activities_doc():
    doc = Document()
    set_style(doc)
    add_cover(doc, 'Verus.AI', 'Relatorio de Atividades\nHomologacao V2', '2.0', '31/03/2026')

    doc.add_heading('1. Resumo Executivo', level=1)
    doc.add_paragraph(
        'Este documento descreve as atividades realizadas no sistema Verus.AI entre a versao 1.0 '
        '(02/03/2026) e a versao 2.0 (31/03/2026), para o cliente PRODERJ. As evolucoes abrangem '
        'novos tipos de documentos, importacao cross-blueprint, ferramentas de pesquisa de mercado '
        'integradas a agentes de IA, melhorias de infraestrutura e correcoes criticas.'
    )

    doc.add_heading('2. Novos Tipos de Documento', level=1)

    doc.add_heading('2.1. Minuta de Edital', level=2)
    doc.add_paragraph('Criacao completa do blueprint "Minuta de Edital" com 10 secoes, 10 agentes geradores e 10 agentes validadores. Cada agente possui system_prompt e user_prompt_template especializados para a secao correspondente, seguindo a Lei 14.133/2021.')
    add_table(doc,
        ['Item', 'Quantidade'],
        [
            ['Secoes', '10'],
            ['Agentes Geradores', '10'],
            ['Agentes Validadores', '10'],
            ['Seed', 'seed_edital_agents.py'],
        ]
    )

    doc.add_heading('2.2. Minuta de Contrato', level=2)
    doc.add_paragraph('Criacao do blueprint "Minuta de Contrato" com 6 secoes, 6 geradores e 6 validadores. Inclui SectionImportConfig para importacao TR -> Contrato.')
    add_table(doc,
        ['Item', 'Quantidade'],
        [
            ['Secoes', '6'],
            ['Agentes Geradores', '6'],
            ['Agentes Validadores', '6'],
            ['Seed', 'seed_minuta_contrato_agents.py'],
        ]
    )

    doc.add_heading('3. Importacao Cross-Blueprint (DOD -> ETP -> TR -> Minuta)', level=1)

    doc.add_heading('3.1. DOD -> ETP (Secao 3 -> Secao 1)', level=2)
    doc.add_paragraph(
        'Implementacao completa do fluxo de importacao de conteudo do DOD (Secao 3 - Demanda) '
        'para o ETP (Secao 1 - Descricao da Necessidade). O usuario seleciona um DOD ja gerado '
        'e o sistema pre-preenche o objetivo do ETP com o conteudo importado.'
    )
    doc.add_paragraph('Arquivos criados/modificados:', style='List Bullet')
    doc.add_paragraph('Backend: etp_import_service.py, etp_views.py, urls.py', style='List Bullet 2')
    doc.add_paragraph('Frontend: use-etp-import.ts, SessionSidebar.tsx, etp/page.tsx', style='List Bullet 2')
    doc.add_paragraph('Seed: seed_etp_rj_section1_dod.py', style='List Bullet 2')

    doc.add_heading('3.2. ETP -> TR', level=2)
    doc.add_paragraph('Importacao ja existente na V1. Mantida e estabilizada.')

    doc.add_heading('3.3. TR -> Minutas', level=2)
    doc.add_paragraph('Implementacao do MinutaImportService para importar conteudo do TR tanto para Minuta de Edital quanto Minuta de Contrato, via SectionImportConfig.')

    doc.add_heading('4. Agent Tools (Ferramentas de Agentes)', level=1)

    doc.add_heading('4.1. Arquitetura', level=2)
    doc.add_paragraph(
        'Criacao de sistema de ferramentas externas para agentes de IA. As ferramentas sao '
        'executadas ANTES da chamada ao LLM (pre-execution pattern), tornando o sistema compativel '
        'com qualquer provider LLM (watsonx, Claude, GPT).'
    )
    doc.add_paragraph('Arquivos criados:', style='List Bullet')
    doc.add_paragraph('models/agent_tools.py - AgentTool e AgentToolLink', style='List Bullet 2')
    doc.add_paragraph('services/agent_tools_service.py - Orquestrador', style='List Bullet 2')
    doc.add_paragraph('services/tools/web_search_tool.py - Busca Google via Serper', style='List Bullet 2')
    doc.add_paragraph('services/tools/pncp_search_tool.py - Busca PNCP com detalhes de itens', style='List Bullet 2')

    doc.add_heading('4.2. Integracao com ETP Secoes 5 e 6', level=2)
    doc.add_paragraph(
        'Os agentes das secoes 5 (Pesquisa de Precos) e 6 (Estimativa de Quantidades e Precos) '
        'do ETP RJ foram atualizados com prompts especializados que utilizam dados reais de '
        'mercado obtidos via Agent Tools. Implementacao de mini-DAG: secao 5 executa tools e '
        'salva resultados no estado; secao 6 herda via tool_results.'
    )

    doc.add_heading('5. Pesquisa de Precos - Serper.dev', level=1)
    doc.add_paragraph(
        'Substituicao do DuckDuckGo pelo Serper.dev (Google Search API) para busca de precos '
        'de mercado. Custo: US$ 0.001 por busca. Resultados muito superiores em qualidade e '
        'relevancia para licitacoes brasileiras.'
    )
    doc.add_paragraph('Correcao de links PNCP: /compras/ substituido por /editais/ (404 corrigido)', style='List Bullet')
    doc.add_paragraph('Enriquecimento de itens PNCP: busca detalhes de preco por item via API', style='List Bullet')
    doc.add_paragraph('Contexto de busca otimizado: "(pregao OR edital OR licitacao)"', style='List Bullet')

    doc.add_heading('6. Validadores ETP RJ', level=1)
    doc.add_paragraph(
        'Auditoria e correcao de todos os 29 validadores do ETP RJ. Os validadores das secoes '
        '8-29 estavam com mismatch (prompt da secao errada). Recriados a partir de backup JSON.'
    )
    doc.add_paragraph('Seed: seed_etp_rj_validators.py (29 validadores corrigidos)', style='List Bullet')

    doc.add_heading('7. Infraestrutura', level=1)

    doc.add_heading('7.1. Cloudflare R2 (Storage)', level=2)
    doc.add_paragraph(
        'Correcao critica: Django 5.2 requer STORAGES dict (nao DEFAULT_FILE_STORAGE string). '
        'Upload de PDF/DOCX para R2 nao estava funcionando. Corrigido em local.py e production.py.'
    )

    doc.add_heading('7.2. Celery Tasks para Upload R2', level=2)
    doc.add_paragraph(
        'Migracao de uploads R2 de threading para Celery tasks assincronas com retry automatico '
        '(3 tentativas, delay de 10s). Tasks: save_pdf_to_r2_task e save_docx_to_r2_task.'
    )

    doc.add_heading('7.3. Objective Summary (Limite de Embedding)', level=2)
    doc.add_paragraph(
        'O modelo de embedding E5-Large tem limite de 512 tokens. Objetivos longos causavam erro. '
        'Solucao: LLM resume o objetivo para max 200 palavras antes da query RAG. Campo '
        'objective_summary adicionado ao StateGraph TypedDict.'
    )

    doc.add_heading('8. Melhorias de Frontend', level=1)

    doc.add_heading('8.1. ReactMarkdown na Analise de Editais', level=2)
    doc.add_paragraph('Modal de analise migrado de parsing manual de linhas para ReactMarkdown + rehypeRaw + remarkGfm. Suporta tabelas, listas, formatacao completa. Modal aumentado para max-w-5xl.')

    doc.add_heading('8.2. Copilot - Textarea e Mensagens', level=2)
    doc.add_paragraph('Textarea de edicao com auto-expansao baseada em scrollHeight. Mensagens com largura max 90% (w-full + max-w-[90%]) para evitar recolhimento.')

    doc.add_heading('8.3. Sub-secoes - Campos Estruturados', level=2)
    doc.add_paragraph('Correcao: campos {{input}} de sub-secoes nao chegavam ao prompt do agente. Adicionadas variaveis input, user_input, fields_data no DynamicGraphBuilder. Criado path should_use_fields_only para sub-secoes com formulario mas sem agente.')

    doc.add_heading('8.4. Scores sem Dizima', level=2)
    doc.add_paragraph('Todos os validation_score arredondados com round(..., 2) antes de salvar no banco. Afeta: base_validator.py (final_score, structural_score, semantic_score) e dynamic_graph_builder.py (score, average_score).')

    doc.add_heading('9. Analise de Editais Configuravel', level=1)
    doc.add_paragraph(
        'Criacao do modelo AIAnalysisConfig no Django Admin (app core) para configurar a analise '
        'de editais. Permite definir provider/modelo LLM, system_prompt personalizado, perguntas '
        'de analise (JSONField), temperatura e max_tokens. Antes era hardcoded para Claude Haiku.'
    )

    doc.add_heading('10. Seeds de Producao', level=1)
    doc.add_paragraph('Lista de management commands criados para aplicar configuracoes em producao:')
    add_table(doc,
        ['Seed', 'Funcao'],
        [
            ['seed_edital_agents', '10 geradores + 10 validadores para Minuta de Edital'],
            ['seed_minuta_contrato_agents', '6 geradores + 6 validadores para Minuta de Contrato'],
            ['seed_etp_rj_validators', '29 validadores ETP RJ corrigidos (de backup JSON)'],
            ['seed_etp_rj_tools_agents', 'Prompts secoes 5 e 6 com Agent Tools'],
            ['seed_etp_rj_section1_dod', 'Prompt secao 1 com importacao DOD + SectionImportConfig'],
            ['seed_agent_tools', 'Web Search (Serper) e PNCP Search tools'],
        ]
    )

    doc.add_heading('11. Correcoes Criticas', level=1)
    add_table(doc,
        ['Bug', 'Causa', 'Correcao'],
        [
            ['Embedding 512 token error', 'objective_summary nao existia no TypedDict', 'Adicionado ao StateGraph + LLM resume objetivo'],
            ['Validadores usando Anthropic', 'Agentes configurados com provider errado', 'SQL UPDATE para watsonx em producao'],
            ['R2 upload nao funcionava', 'DEFAULT_FILE_STORAGE ignorado no Django 5.2', 'STORAGES dict em settings'],
            ['PNCP links 404', '/compras/ nao e rota publica', 'Substituido por /editais/'],
            ['Secao 6 sempre igual', 'Prompt antigo contraditorio', 'Reescrito com dados reais da secao 5'],
            ['DOD sessions nao encontradas', 'Buscava so GenerationSession (novo)', 'Adicionado fallback IntelligentSession (legado)'],
            ['Sub-secao campos ignorados', 'should_generate exigia generator_agent', 'Adicionado path should_use_fields_only'],
            ['selectedDodSessionId resetado', 'Set null apos criacao de sessao', 'Removido reset'],
            ['Scores com dizima', 'Divisao float sem arredondamento', 'round(..., 2) antes de salvar'],
        ]
    )

    # SALVAR
    path = os.path.join(OUTPUT_DIR, 'verus_ai_Relatorio_Atividades_V2.docx')
    doc.save(path)
    print(f'[OK] {path}')
    return path


# ============================================================
# DOCUMENTO 3: Roadmap de Evolucao
# ============================================================
def generate_roadmap_doc():
    doc = Document()
    set_style(doc)
    add_cover(doc, 'Verus.AI', 'Roadmap de Evolucao\nProximas Entregas', '1.0', '31/03/2026')

    doc.add_heading('Visao Geral do Roadmap', level=1)
    add_table(doc,
        ['#', 'Iniciativa', 'Complexidade', 'Impacto', 'Prioridade'],
        [
            ['1', 'Consolidar KBs em query_by_links()', 'Baixa', 'Tecnico/Admin', 'Curto prazo'],
            ['2', 'Continuidade de Geracao', 'Media', 'Alto (UX)', 'Curto prazo'],
            ['3', 'Extracao Inteligente com Docling', 'Alta', 'Alto (Produto)', 'Medio prazo'],
            ['4', 'GraphRAG com Neo4j', 'Muito Alta', 'Alto (Qualidade IA)', 'Longo prazo'],
        ]
    )
    doc.add_page_break()

    # ITEM 1
    doc.add_heading('1. Consolidar Descoberta de KBs', level=1)

    doc.add_heading('1.1. Problema', level=2)
    doc.add_paragraph(
        'O sistema tem dois mecanismos de descoberta de KBs coexistindo: '
        'M2M direto legado (SectionAgentConfig.knowledge_bases) + metodo query() com 3 camadas, '
        'e o novo AgentKnowledgeBaseLink com priority/purpose/top_k + metodo query_by_links(). '
        'Alem disso, query_by_links() faz fallback invisivel para KBs globais/blueprint via FK, '
        'tornando o admin confuso (links vazios mas agente recebe KBs).'
    )

    doc.add_heading('1.2. Estado Atual (ja implementado)', level=2)
    add_table(doc,
        ['Recurso', 'Status'],
        [
            ['Prioridade de consulta (priority)', 'Implementado'],
            ['Proposito por KB (purpose)', 'Implementado'],
            ['top_k por KB', 'Implementado'],
            ['min_similarity por KB', 'Implementado'],
            ['Filtro de fontes (selected_sources)', 'Implementado'],
            ['Resumo interpretativo (include_summary)', 'Implementado'],
            ['Admin inline no SectionAgentConfig', 'Implementado'],
        ]
    )

    doc.add_heading('1.3. Etapas de Implementacao', level=2)

    doc.add_heading('Etapa 1 - Migrar chamadores de query() para query_by_links()', level=3)
    doc.add_paragraph('Identificar todos os pontos que ainda chamam kb_service.query() (legado). Substituir por kb_service.query_by_links(). Validar equivalencia.')
    doc.add_paragraph('Arquivos: knowledge_base_service.py, qualquer service/view que chame query()')

    doc.add_heading('Etapa 2 - Auto-criar AgentKnowledgeBaseLink ao criar KB de blueprint', level=3)
    doc.add_paragraph('Signal post_save no KnowledgeBase: ao criar KB com kb_layer=blueprint, auto-criar links para todos os agentes do blueprint. Defaults: priority=50, purpose=reference, top_k=5, min_similarity=0.6.')
    doc.add_paragraph('Arquivos: models/knowledge_base.py (signal), admin.py')

    doc.add_heading('Etapa 3 - Remover M2M direto legado', level=3)
    doc.add_paragraph('Remover campo SectionAgentConfig.knowledge_bases (M2M direto). Nova migracao Django.')
    doc.add_paragraph('Arquivos: models/agent.py, admin.py, migracao')

    doc.add_heading('Etapa 4 - Remover metodo query() legado', level=3)
    doc.add_paragraph('Deletar query() de KnowledgeBaseService. Remover imports e referencias.')
    doc.add_paragraph('Arquivos: knowledge_base_service.py')

    doc.add_heading('1.4. Verificacao', level=2)
    doc.add_paragraph('Admin do SectionAgentConfig mostra todos os links explicitos', style='List Bullet')
    doc.add_paragraph('Criar KB de blueprint -> links auto-criados para agentes', style='List Bullet')
    doc.add_paragraph('Geracao funciona identicamente (mesmos resultados RAG)', style='List Bullet')
    doc.add_paragraph('Nenhum chamador usa query() antigo', style='List Bullet')
    doc.add_page_break()

    # ITEM 2
    doc.add_heading('2. Continuidade de Geracao', level=1)

    doc.add_heading('2.1. Problema', level=2)
    doc.add_paragraph(
        'O usuario gera 10 de 30 secoes hoje. Amanha, quer continuar gerando as 20 restantes. '
        'Hoje so existe "Regenerar" (refaz UMA secao) e "Nova Geracao" (comeca do ZERO). '
        'Nao existe "Continuar" - retomar de onde parou.'
    )

    doc.add_heading('2.2. Cenario de Uso', level=2)
    doc.add_paragraph('Dia 1: Usuario cria sessao ETP, seleciona secoes 1-10, gera com sucesso.', style='List Bullet')
    doc.add_paragraph('Dia 2: Abre historico, ve 10/30 secoes, clica "Continuar Geracao", seleciona 11-20, sistema gera APENAS as novas usando as anteriores como contexto.', style='List Bullet')
    doc.add_paragraph('Dia 3: Continua com 21-30. Documento completo.', style='List Bullet')

    doc.add_heading('2.3. Etapas de Implementacao', level=2)

    doc.add_heading('Etapa 1 - Backend: endpoint de continuidade', level=3)
    doc.add_paragraph('Novo endpoint POST /generate/continue/ que carrega GenerationSession existente, identifica secoes completed vs pending, monta DynamicGraphBuilder apenas para as pendentes, usando as completadas como previous_sections.')
    doc.add_paragraph('Arquivos: generation_views.py, dynamic_graph_builder.py')

    doc.add_heading('Etapa 2 - Backend: injetar secoes existentes como contexto', level=3)
    doc.add_paragraph('DynamicGraphBuilder recebe dict pre_existing_sections: {section_number: content}. No _build_previous_sections_context(), incluir conteudo das pre-existentes. NAO criar nos de geracao/validacao para essas secoes.')
    doc.add_paragraph('Arquivos: dynamic_graph_builder.py')

    doc.add_heading('Etapa 3 - Frontend: detectar sessao incompleta', level=3)
    doc.add_paragraph('Ao abrir sessao existente, verificar completed_sections < total_sections. Se incompleta, mostrar badge "10/30 secoes" + botao "Continuar Geracao".')
    doc.add_paragraph('Arquivos: use-intelligent-assistant.ts, GenerationPhase.tsx, SessionSidebar.tsx, etp/page.tsx')

    doc.add_heading('Etapa 4 - Frontend: seletor de secoes para continuar', level=3)
    doc.add_paragraph('Componente que mostra secoes completadas (check verde, desabilitado) e pendentes (checkbox habilitado, pre-selecionado). Botao "Gerar X secoes selecionadas".')

    doc.add_heading('2.4. Verificacao', level=2)
    doc.add_paragraph('Gerar 5 secoes de um ETP com 30', style='List Bullet')
    doc.add_paragraph('Fechar browser, reabrir, abrir historico', style='List Bullet')
    doc.add_paragraph('Sessao mostra "5/30 secoes" + botao continuar', style='List Bullet')
    doc.add_paragraph('Continuar -> gera secoes 6-15 usando 1-5 como contexto', style='List Bullet')
    doc.add_paragraph('Secoes 1-5 NAO foram regeneradas', style='List Bullet')
    doc.add_paragraph('Secao 6 faz referencia ao conteudo da secao 5 (contexto preservado)', style='List Bullet')
    doc.add_page_break()

    # ITEM 3
    doc.add_heading('3. Extracao Inteligente de Documento Fonte', level=1)

    doc.add_heading('3.1. Problema', level=2)
    doc.add_paragraph(
        'O usuario ja tem um documento pronto (ex: pre-ETP em Word/PDF) e quer usa-lo como base '
        'para gerar o ETP oficial. O upload atual so alimenta o RAG - o agente recebe chunks '
        'aleatorios sem saber qual trecho corresponde a qual secao. Resultado: IA "se inspira" '
        'mas nao extrai fielmente.'
    )

    doc.add_heading('3.2. Solucao: Docling + LLM Extrator', level=2)
    doc.add_paragraph(
        'IBM Docling (open source) faz parsing avancado preservando estrutura hierarquica. '
        'LLM extrator mapeia trechos do documento para secoes do blueprint alvo.'
    )

    doc.add_heading('Por que Docling', level=3)
    add_table(doc,
        ['Aspecto', 'pypdf/python-docx (atual)', 'Docling (IBM Research)'],
        [
            ['Tabelas complexas', 'Perde estrutura', 'Preserva linhas/colunas'],
            ['Cabecalhos/rodapes', 'Mistura com conteudo', 'Separa automaticamente'],
            ['Hierarquia (secao/subsecao)', 'Texto plano', 'Arvore estruturada'],
            ['OCR (PDFs escaneados)', 'Nao suporta', 'Suporta nativamente'],
            ['Numeracao inconsistente', 'Texto plano', 'Detecta padroes'],
        ]
    )

    doc.add_heading('3.3. Etapas de Implementacao', level=2)

    doc.add_heading('Fase 1 - Docling Service (parsing avancado)', level=3)
    doc.add_paragraph('Novo service docling_service.py que recebe arquivo, processa com Docling, extrai estrutura hierarquica (secoes, paragrafos, tabelas), retorna documento estruturado.')
    doc.add_paragraph('Dependencia: pip install docling')

    doc.add_heading('Fase 2 - LLM Extrator (mapeamento secao-a-secao)', level=3)
    doc.add_paragraph('Novo agente que recebe documento estruturado + lista de secoes do blueprint. Retorna mapeamento: {section_number: extracted_content, confidence}. Confidence score 0-1 por mapeamento.')

    doc.add_heading('Fase 3 - API de Extracao', level=3)
    doc.add_paragraph('POST /extract/ (upload + blueprint_id) -> processamento assincrono via Celery. GET /extract/{id}/ -> resultado com mapeamentos e stats.')

    doc.add_heading('Fase 4 - Frontend: revisao de extracao', level=3)
    doc.add_paragraph('Tela mostra "Extraimos conteudo para 15/30 secoes" com badges de confianca (verde >= 80%, amarelo >= 50%, vermelho < 50%). Usuario revisa/edita cada extracao. Secoes confirmadas viram imported_content.')

    doc.add_heading('3.4. Verificacao', level=2)
    doc.add_paragraph('Upload de pre-ETP real do PRODERJ', style='List Bullet')
    doc.add_paragraph('Docling extrai secoes estruturadas', style='List Bullet')
    doc.add_paragraph('LLM mapeia para secoes do blueprint', style='List Bullet')
    doc.add_paragraph('Frontend mostra revisao com badges', style='List Bullet')
    doc.add_paragraph('Geracao usa imported_content + RAG normal', style='List Bullet')
    doc.add_page_break()

    # ITEM 4
    doc.add_heading('4. GraphRAG com Neo4j', level=1)

    doc.add_heading('4.1. Problema', level=2)
    doc.add_paragraph(
        'O RAG atual (pgvector) e flat - busca por similaridade vetorial pura. Nao entende '
        'relacoes entre entidades, nao conecta mencoes ao mesmo orgao/lei em diferentes chunks, '
        'e nao permite queries estruturais como "quais artigos da Lei 14.133 regulam pregao".'
    )

    doc.add_heading('4.2. Arquitetura Proposta', level=2)
    doc.add_paragraph(
        'Hybrid retrieval: pgvector (busca vetorial por similaridade) + Neo4j (traversal por '
        'relacoes semanticas). Combinacao dos dois para contexto mais rico.'
    )

    doc.add_heading('Entidades', level=3)
    add_table(doc,
        ['Entidade', 'Exemplos'],
        [
            ['Orgao', 'PRODERJ, SEFAZ-RJ, TCE-RJ'],
            ['Lei', 'Lei 14.133/2021, Decreto 10.541/2024'],
            ['Artigo', 'Art. 18 da Lei 14.133'],
            ['Modalidade', 'Pregao Eletronico, Concorrencia'],
            ['Objeto', 'Solucao de TI para gestao documental'],
            ['Valor', 'R$ 1.500.000,00'],
            ['Empresa', 'Fornecedor X, CNPJ Y'],
        ]
    )

    doc.add_heading('Relacoes', level=3)
    add_table(doc,
        ['Relacao', 'Exemplo'],
        [
            ['REGULAMENTA', 'Lei 14.133 -> Pregao Eletronico'],
            ['CONTIDO_EM', 'Art. 18 -> Lei 14.133'],
            ['PUBLICADO_POR', 'ETP -> PRODERJ'],
            ['REFERE_VALOR', 'Contrato -> R$ 1.5M'],
            ['FORNECE', 'Empresa X -> Solucao de TI'],
            ['PRECEDE', 'DOD -> ETP -> TR'],
        ]
    )

    doc.add_heading('4.3. Fases de Implementacao', level=2)

    doc.add_heading('Fase 1 - Infraestrutura + Extracao de Entidades', level=3)
    doc.add_paragraph('Container Neo4j no docker-compose. Neo4j Service (graph_rag_service.py) para CRUD de entidades. Pipeline de extracao: LLM extrai entidades/relacoes dos chunks no upload.')
    doc.add_paragraph('Dependencias: neo4j (Python driver), Neo4j Community 5.x, APOC plugin')

    doc.add_heading('Fase 2 - Hybrid Retrieval (pgvector + Cypher)', level=3)
    doc.add_paragraph('Query Rewriter: LLM analisa query e decide se precisa busca vetorial, grafo ou ambas. Se grafo, gera query Cypher. Hybrid Ranker funde resultados: final_score = alpha * vector_score + beta * graph_score.')
    doc.add_paragraph('Novo metodo query_hybrid() no KnowledgeBaseService.')

    doc.add_heading('Fase 3 - Agent Tools com Acesso ao Grafo', level=3)
    doc.add_paragraph('Novo AgentTool tipo graph_query. Agente gera Cypher dinamico baseado na necessidade da secao. Resultado injetado no contexto (mesmo padrao Agent Tools).')
    doc.add_paragraph('Exemplo: Agente da secao 5 busca "empresas que forneceram solucao de TI com valores" via Cypher.')

    doc.add_heading('4.4. Verificacao', level=2)
    doc.add_paragraph('Fase 1: Upload documento -> Neo4j Browser mostra entidades criadas', style='List Bullet')
    doc.add_paragraph('Fase 2: Geracao ETP -> log mostra pgvector + Neo4j resultados combinados', style='List Bullet')
    doc.add_paragraph('Fase 3: Agente secao 5 usa tool graph_query -> retorna dados reais do grafo', style='List Bullet')

    # SALVAR
    path = os.path.join(OUTPUT_DIR, 'verus_ai_Roadmap_Evolucao.docx')
    doc.save(path)
    print(f'[OK] {path}')
    return path


if __name__ == '__main__':
    print('Gerando documentos DOCX...\n')
    generate_system_doc()
    generate_activities_doc()
    generate_roadmap_doc()
    print('\nConcluido! Documentos salvos em docs/')
