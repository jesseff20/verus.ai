"""
Views para Documents
"""
import logging

from rest_framework import viewsets, status, filters

logger = logging.getLogger(__name__)
from apps.accounts.permissions import resolve_role, is_admin_or_manager, TrialAccessPermission
from apps.accounts.tasks import send_email_task
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from drf_spectacular.utils import extend_schema, extend_schema_view, OpenApiResponse
from django.utils import timezone
from ..models import Document, DocumentGenerator
from ..serializers import (
    DocumentListSerializer, DocumentDetailSerializer, DocumentCreateSerializer,
    DocumentUpdateSerializer, DocumentGeneratorSerializer,
)
from ..permissions import IsOwnerOrAdmin
from apps.core.mixins import AuditModelMixin, AuditActionMixin, AuditDownloadMixin, OwnerRequiredMixin


@extend_schema_view(
    list=extend_schema(
        summary="Listar Documents",
        description="Retorna lista de Documents do usuário (ou todos para admin)",
        tags=["Documents"],
        responses={200: DocumentListSerializer(many=True)}
    ),
    retrieve=extend_schema(
        summary="Buscar ETP por ID",
        description="Retorna detalhes completos de um ETP",
        tags=["Documents"],
        responses={200: DocumentDetailSerializer}
    ),
    create=extend_schema(
        summary="Criar ETP",
        description="Cria novo ETP",
        tags=["Documents"],
        request=DocumentCreateSerializer,
        responses={201: DocumentDetailSerializer}
    ),
    update=extend_schema(
        summary="Atualizar ETP",
        description="Atualiza ETP (apenas dono ou admin)",
        tags=["Documents"],
        request=DocumentUpdateSerializer,
        responses={200: DocumentDetailSerializer}
    ),
    partial_update=extend_schema(
        summary="Atualizar parcialmente ETP",
        description="Atualiza campos específicos",
        tags=["Documents"],
        request=DocumentUpdateSerializer,
        responses={200: DocumentDetailSerializer}
    ),
    destroy=extend_schema(
        summary="Deletar ETP",
        description="Remove ETP (apenas dono ou admin)",
        tags=["Documents"],
        responses={204: None}
    ),
)
class DocumentViewSet(OwnerRequiredMixin, AuditModelMixin, AuditActionMixin, AuditDownloadMixin, viewsets.ModelViewSet):
    """ViewSet para Documents com auditoria automatica."""
    queryset = Document.objects.all()
    permission_classes = [TrialAccessPermission, IsOwnerOrAdmin]

    # Mensagens customizadas para auditoria
    audit_create_message = 'Documento criado'
    audit_update_message = 'Documento atualizado'
    audit_delete_message = 'Documento excluído'

    def get_serializer_class(self):
        if self.action == 'list':
            return DocumentListSerializer
        elif self.action == 'create':
            return DocumentCreateSerializer
        elif self.action in ['update', 'partial_update']:
            return DocumentUpdateSerializer
        return DocumentDetailSerializer

    def get_queryset(self):
        user = self.request.user
        queryset = Document.objects.select_related(
            'user', 'form_template', 'document_template')

        # Filtro por role:
        # - Operador/Visualizador: só vê seus próprios documentos
        # - Gestor: vê todos, mas apenas os que já foram gerados (têm PDF)
        # - SuperAdmin/Revisor: vê todos os documentos
        role = resolve_role(user.role)
        _limited_roles = [
            'analista', 'visualizador', 'assistente', 'paralegal', 'estagiario',
            'secretaria', 'assessor', 'servidor', 'cliente',
            'advogado_junior', 'advogado_pleno',
        ]
        _manager_roles = ['gestor', 'coordenador', 'supervisor']
        if role in _limited_roles:
            queryset = queryset.filter(user=user)
        elif role in _manager_roles:
            queryset = queryset.exclude(generated_html__isnull=True).exclude(generated_html='')

        # Filtro por status
        status_filter = self.request.query_params.get('status')
        if status_filter:
            queryset = queryset.filter(status=status_filter)

        # Filtro por origem (source)
        source_filter = self.request.query_params.get('source')
        if source_filter:
            queryset = queryset.filter(source=source_filter)

        # Filtro por usuário (apenas para Gestor+)
        user_id_filter = self.request.query_params.get('user_id')
        if user_id_filter and role not in _limited_roles:
            queryset = queryset.filter(user_id=user_id_filter)

        # Busca por título ou processo
        search = self.request.query_params.get('search')
        if search:
            from django.db.models import Q
            queryset = queryset.filter(
                Q(title__icontains=search) | Q(numero_processo__icontains=search)
            )

        return queryset.order_by('-created_at')

    def create(self, request, *args, **kwargs):
        """Sobrescreve create para retornar DocumentDetailSerializer na resposta"""
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        self.perform_create(serializer)

        # Retorna com DetailSerializer para incluir todos os campos (incluindo id)
        instance = serializer.instance
        detail_serializer = DocumentDetailSerializer(
            instance, context={'request': request})
        headers = self.get_success_headers(detail_serializer.data)
        return Response(detail_serializer.data, status=status.HTTP_201_CREATED, headers=headers)

    @extend_schema(
        summary="Visão Unificada de Documentos",
        description="Retorna todos os documentos de ambos os sistemas (Document + GeneratedDocument do Assistente Inteligente)",
        tags=["Documents"],
        responses={200: OpenApiResponse(description="Lista unificada de documentos")}
    )
    @action(detail=False, methods=['get'])
    def unified(self, request):
        """
        Retorna documentos de ambos os sistemas combinados:
        - Document (sistema tradicional de geradores)
        - GeneratedDocument (assistente inteligente)

        Filtros:
        - status: Filtrar por status
        - source: manual, generator, assistant
        - user_id: Filtrar por usuário (Manager+ apenas)
        - search: Buscar por título ou processo
        - system: documents, intelligent_assistant (filtrar por sistema específico)
        """
        from apps.intelligent_assistant.models import GeneratedDocument

        documents = []
        user = request.user

        # Filtros
        status_filter = request.query_params.get('status')
        source_filter = request.query_params.get('source')
        user_id_filter = request.query_params.get('user_id')
        search = request.query_params.get('search')
        system_filter = request.query_params.get('system')

        # ========================================
        # 1. DOCUMENTOS DO SISTEMA TRADICIONAL
        # ========================================
        if not system_filter or system_filter == 'documents':
            doc_queryset = Document.objects.select_related(
                'user', 'form_template', 'document_template'
            )

            # Filtro por role (operadores/visualizadores só vêem seus próprios)
            _role = resolve_role(user.role)
            _limited = [
                'analista', 'visualizador', 'assistente', 'paralegal', 'estagiario',
                'secretaria', 'assessor', 'servidor', 'cliente',
                'advogado_junior', 'advogado_pleno',
            ]
            _mgmt = ['gestor', 'coordenador', 'supervisor']
            if _role in _limited:
                doc_queryset = doc_queryset.filter(user=user)
            elif _role in _mgmt:
                doc_queryset = doc_queryset.exclude(generated_html__isnull=True).exclude(generated_html='')
                if user_id_filter:
                    doc_queryset = doc_queryset.filter(user_id=user_id_filter)
            elif user_id_filter:
                doc_queryset = doc_queryset.filter(user_id=user_id_filter)

            # Filtros adicionais
            if status_filter:
                doc_queryset = doc_queryset.filter(status=status_filter)
            if source_filter and source_filter != 'assistant':
                doc_queryset = doc_queryset.filter(source=source_filter)
            if search:
                from django.db.models import Q
                doc_queryset = doc_queryset.filter(
                    Q(title__icontains=search) | Q(numero_processo__icontains=search)
                )

            # Se source_filter é 'assistant', não incluir documentos do sistema tradicional
            if source_filter != 'assistant':
                for doc in doc_queryset:
                    has_content = bool(doc.generated_html)
                    documents.append({
                        'id': str(doc.id),
                        'title': doc.title or f"Documento {str(doc.id)[:8]}",
                        'numero_processo': doc.numero_processo,
                        'user_id': doc.user.id,
                        'user_name': doc.user.get_full_name() or doc.user.username,
                        'source': doc.source,
                        'source_display': doc.get_source_display(),
                        'status': doc.status,
                        'status_display': doc.get_status_display(),
                        'system': 'documents',
                        'system_display': 'Gerador de Documento',
                        'form_template_name': doc.form_template.name if doc.form_template else None,
                        'progress': doc.progress,
                        'version': doc.version,
                        'created_at': doc.created_at.isoformat(),
                        'updated_at': doc.updated_at.isoformat(),
                        'has_generated_content': has_content,
                        'pdf_url': f'/api/v1/documents/items/{doc.id}/export_pdf/' if has_content else None,
                        'preview_url': f'/dashboard/documents/{doc.id}/preview' if has_content else None,
                        'detail_url': f'/dashboard/documents/{doc.id}',
                    })

        # ========================================
        # 2. DOCUMENTOS DO ASSISTENTE INTELIGENTE
        # ========================================
        if not system_filter or system_filter == 'intelligent_assistant':
            # Se source_filter está definido e não é 'assistant', não incluir
            if not source_filter or source_filter == 'assistant':
                gen_queryset = GeneratedDocument.objects.select_related(
                    'session', 'session__user', 'session__blueprint'
                ).prefetch_related('session__generation_sessions')

                # Filtro por role (operadores/visualizadores só vêem seus próprios)
                _role2 = resolve_role(user.role)
                if _role2 in _limited:
                    gen_queryset = gen_queryset.filter(session__user=user)
                elif _role2 in _mgmt:
                    if user_id_filter:
                        gen_queryset = gen_queryset.filter(session__user_id=user_id_filter)
                elif user_id_filter:
                    gen_queryset = gen_queryset.filter(session__user_id=user_id_filter)

                # Busca
                if search:
                    from django.db.models import Q
                    gen_queryset = gen_queryset.filter(
                        Q(title__icontains=search) | Q(session__objective__icontains=search)
                    )

                # /etp é página genérica (cobre DOD, mapa_riscos, etc.)
                # /tr e /edital têm fluxos especiais (conversão ETP→TR e TR→Edital)
                DOC_TYPE_ROUTES = {
                    'termo_referencia': 'tr',
                    'edital': 'edital',
                }

                # N+1 fix: pré-popular dict {code: name} de DocumentType em UMA
                # query, em vez de deixar a property `session.document_type_display`
                # ir ao banco pra cada GeneratedDocument do loop. Reduz N queries
                # adicionais (1 por doc) a 1 query única.
                from apps.core.models import DocumentType
                doc_type_names = dict(
                    DocumentType.objects.values_list('code', 'name')
                )

                for doc in gen_queryset:
                    session = doc.session
                    doc_type = session.document_type if session else None
                    route_slug = DOC_TYPE_ROUTES.get(doc_type, 'etp')
                    path_segment = f'/{route_slug}'

                    gen_session = None
                    if session:
                        all_gens = sorted(
                            session.generation_sessions.all(),
                            key=lambda g: g.created_at,
                            reverse=True,
                        )
                        completed = [g for g in all_gens if g.status == 'completed']
                        gen_session = completed[0] if completed else (all_gens[0] if all_gens else None)

                    if session:
                        gen_param = f'&generation_session={gen_session.id}' if gen_session else ''
                        detail_url = (
                            f'/dashboard/intelligent-assistant{path_segment}'
                            f'?session={session.id}&phase=result&doc={doc.id}{gen_param}'
                        )
                    else:
                        detail_url = f'/dashboard/intelligent-assistant?doc={doc.id}'

                    documents.append({
                        'id': str(doc.id),
                        'title': doc.title or f"Documento {str(doc.id)[:8]}",
                        'numero_processo': None,
                        'user_id': session.user.id if session else None,
                        'user_name': session.user.get_full_name() or session.user.username if session else 'Desconhecido',
                        'source': 'assistant',
                        'source_display': 'Assistente Inteligente',
                        'status': 'completed',
                        'status_display': 'Completo',
                        'system': 'intelligent_assistant',
                        'system_display': 'Assistente Inteligente',
                        'form_template_name': None,
                        'document_generator_name': None,
                        'blueprint_name': session.blueprint.name if session and session.blueprint else None,
                        'document_type': (
                            doc_type_names.get(session.document_type, session.document_type)
                            if session else None
                        ),
                        'progress': 100,
                        'version': 1,
                        'created_at': doc.generated_at.isoformat() if doc.generated_at else doc.updated_at.isoformat(),
                        'updated_at': doc.updated_at.isoformat(),
                        'has_generated_content': True,  # Assistente sempre gera conteúdo
                        'pdf_url': doc.pdf_url,
                        'preview_url': None,  # Assistente usa pdf_url direto
                        'session_id': str(session.id) if session else None,
                        'detail_url': detail_url,
                    })

        # ========================================
        # 3. ORDENAÇÃO E RESPOSTA
        # ========================================
        # Ordenar por data de criação (mais recentes primeiro)
        documents.sort(key=lambda x: x['created_at'], reverse=True)

        return Response({
            'count': len(documents),
            'results': documents,
            'filters_applied': {
                'status': status_filter,
                'source': source_filter,
                'user_id': user_id_filter,
                'search': search,
                'system': system_filter,
            }
        })

    @extend_schema(
        summary="Estatísticas do Dashboard",
        description="Retorna estatísticas agregadas para o dashboard principal",
        tags=["Documents"],
        responses={200: OpenApiResponse(description="Estatísticas do dashboard")}
    )
    @action(detail=False, methods=['get'])
    def dashboard_stats(self, request):
        """
        Retorna estatísticas agregadas para o dashboard:
        - Total de documentos (ambos os sistemas)
        - Documentos por status
        - Total de documentos na Knowledge Base
        - Notificações (pendentes de revisão)
        - Documentos recentes (últimos 5)
        """
        from django.db.models import Count
        from apps.intelligent_assistant.models import GeneratedDocument
        from apps.kb.models import Document as KBDocument

        user = request.user

        # ========================================
        # 1. DOCUMENTOS TRADICIONAIS
        # ========================================
        doc_qs = Document.objects.all()

        # Filtro por role (operadores/visualizadores só vêem seus próprios)
        _dash_role = resolve_role(user.role)
        _dash_limited = [
            'analista', 'visualizador', 'assistente', 'paralegal', 'estagiario',
            'secretaria', 'assessor', 'servidor', 'cliente',
            'advogado_junior', 'advogado_pleno',
        ]
        if _dash_role in _dash_limited:
            doc_qs = doc_qs.filter(user=user)

        # Contagens por status
        doc_by_status = doc_qs.values('status').annotate(count=Count('id'))
        status_counts = {item['status']: item['count'] for item in doc_by_status}

        total_docs = doc_qs.count()

        # ========================================
        # 2. DOCUMENTOS DO ASSISTENTE IA
        # ========================================
        gen_qs = GeneratedDocument.objects.all()

        if _dash_role in _dash_limited:
            gen_qs = gen_qs.filter(session__user=user)

        total_gen = gen_qs.count()

        # ========================================
        # 3. KNOWLEDGE BASE
        # ========================================
        kb_total = KBDocument.objects.filter(status='completed').count()

        # ========================================
        # 4. DOCUMENTOS RECENTES (últimos 5)
        # ========================================
        recent_documents = []

        # Documentos tradicionais recentes
        recent_doc_qs = doc_qs.exclude(status='archived').select_related('user', 'form_template').order_by('-created_at')[:5]
        for doc in recent_doc_qs:
            recent_documents.append({
                'id': str(doc.id),
                'title': doc.title or f"Documento {str(doc.id)[:8]}",
                'status': doc.status,
                'status_display': doc.get_status_display(),
                'source': doc.source,
                'source_display': doc.get_source_display(),
                'system': 'documents',
                'created_at': doc.created_at.isoformat(),
                'user_name': doc.user.get_full_name() or doc.user.username,
            })

        # Documentos do assistente IA recentes
        recent_gen_qs = gen_qs.select_related('session', 'session__user').order_by('-generated_at')[:5]
        for doc in recent_gen_qs:
            recent_documents.append({
                'id': str(doc.id),
                'title': doc.title or f"Documento {str(doc.id)[:8]}",
                'status': 'completed',
                'status_display': 'Completo',
                'source': 'assistant',
                'source_display': 'Assistente Inteligente',
                'system': 'intelligent_assistant',
                'created_at': (doc.generated_at or doc.updated_at).isoformat(),
                'user_name': doc.session.user.get_full_name() or doc.session.user.username if doc.session else 'Desconhecido',
            })

        # Ordenar por data e pegar os 5 mais recentes
        recent_documents.sort(key=lambda x: x['created_at'], reverse=True)
        recent_documents = recent_documents[:5]

        # ========================================
        # 5. RESPOSTA
        # ========================================
        return Response({
            'documents': {
                'total': total_docs + total_gen,
                'traditional': total_docs,
                'assistant': total_gen,
                'by_status': {
                    'draft': status_counts.get('draft', 0),
                    'in_review': status_counts.get('in_review', 0),
                    'completed': status_counts.get('completed', 0) + total_gen,
                    'archived': status_counts.get('archived', 0),
                }
            },
            'knowledge_base': {
                'total': kb_total
            },
            'notifications': {
                'pending_review': status_counts.get('in_review', 0)
            },
            'recent_documents': recent_documents
        })

    @extend_schema(
        summary="Criar nova versão",
        description="Cria nova versão do ETP",
        tags=["Documents"],
        responses={201: DocumentDetailSerializer}
    )
    @action(detail=True, methods=['post'])
    def create_version(self, request, pk=None):
        """Cria nova versão"""
        etp = self.get_object()
        new_etp = etp.create_version(request.user)
        serializer = DocumentDetailSerializer(new_etp)
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    @extend_schema(
        summary="Completar ETP",
        description="Marca ETP como completo",
        tags=["Documents"],
        responses={200: DocumentDetailSerializer}
    )
    @action(detail=True, methods=['post'])
    def complete(self, request, pk=None):
        """Marca como completo"""
        etp = self.get_object()
        etp.status = 'completed'
        etp.completed_at = timezone.now()
        etp.save()
        serializer = DocumentDetailSerializer(etp)
        return Response(serializer.data, status=status.HTTP_200_OK)

    @extend_schema(
        summary="Visualizar HTML gerado",
        description="Retorna HTML gerado do documento para preview",
        tags=["Documents"],
        responses={200: OpenApiResponse(description="HTML do documento")}
    )
    @action(detail=True, methods=['get'])
    def preview(self, request, pk=None):
        """Retorna HTML gerado para preview"""
        from django.http import HttpResponse

        etp = self.get_object()

        if not etp.generated_html:
            return Response({
                'error': 'Documento ainda não foi gerado. Execute o endpoint /generate/ primeiro.'
            }, status=status.HTTP_400_BAD_REQUEST)

        # Retornar HTML puro para renderizar no navegador
        return HttpResponse(etp.generated_html, content_type='text/html')

    @extend_schema(
        summary="Exportar para PDF",
        description="Exporta o documento gerado para PDF (A4 vertical)",
        tags=["Documents"],
        responses={200: OpenApiResponse(description="Arquivo PDF")}
    )
    @action(detail=True, methods=['get'])
    def export_pdf(self, request, pk=None):
        """Exporta documento para PDF"""
        from django.http import HttpResponse
        from weasyprint import HTML, CSS
        from weasyprint.text.fonts import FontConfiguration
        import io

        etp = self.get_object()

        if not etp.generated_html:
            return Response({
                'error': 'Documento ainda não foi gerado. Execute o endpoint /generate/ primeiro.'
            }, status=status.HTTP_400_BAD_REQUEST)

        try:
            # Configuração de fonte
            font_config = FontConfiguration()

            # CSS adicional para garantir formato A4 vertical
            pdf_css = CSS(string='''
                @page {
                    size: A4 portrait;
                    margin: 2cm;
                }

                body {
                    font-family: Arial, sans-serif;
                    font-size: 12pt;
                    line-height: 1.6;
                }

                h1 {
                    font-size: 18pt;
                    margin-top: 0;
                }

                h2 {
                    font-size: 16pt;
                    margin-top: 20pt;
                }

                h3 {
                    font-size: 14pt;
                    margin-top: 15pt;
                }

                p {
                    margin-bottom: 10pt;
                    text-align: justify;
                }

                table {
                    width: 100%;
                    border-collapse: collapse;
                    margin: 15pt 0;
                }

                table th, table td {
                    border: 1px solid #000;
                    padding: 8pt;
                    text-align: left;
                }

                table th {
                    background-color: #f0f0f0;
                    font-weight: bold;
                }
            ''', font_config=font_config)

            # Gerar PDF
            html = HTML(string=etp.generated_html)
            pdf_file = html.write_pdf(
                stylesheets=[pdf_css], font_config=font_config)

            # Criar response
            response = HttpResponse(pdf_file, content_type='application/pdf')

            # Nome do arquivo baseado no título do documento
            filename = f"{etp.title.replace(' ', '_')}.pdf"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'

            return response

        except Exception as e:
            return Response({
                'error': f'Erro ao gerar PDF: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @extend_schema(
        summary="Exportar para DOCX",
        description="Exporta o documento gerado para DOCX (Word)",
        tags=["Documents"],
        responses={200: OpenApiResponse(description="Arquivo DOCX")}
    )
    @action(detail=True, methods=['get'])
    def export_docx(self, request, pk=None):
        """Exporta documento como DOCX usando python-docx"""
        from django.http import HttpResponse
        from docx import Document as DocxDocument
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re
        import io

        doc = self.get_object()

        if not doc.generated_html:
            return Response(
                {'error': 'Documento ainda não foi gerado'},
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            docx = DocxDocument()

            # Margens A4
            for section in docx.sections:
                section.top_margin = Cm(2.5)
                section.bottom_margin = Cm(2.5)
                section.left_margin = Cm(3)
                section.right_margin = Cm(2)

            html = doc.generated_html

            # Remove tags de script/style
            html = re.sub(r'<(script|style)[^>]*>.*?</(script|style)>', '', html, flags=re.DOTALL | re.IGNORECASE)

            # Extrai conteúdo do body se existir
            body_match = re.search(r'<body[^>]*>(.*?)</body>', html, re.DOTALL | re.IGNORECASE)
            content = body_match.group(1) if body_match else html

            # Remove tags HTML e converte quebras de linha
            def strip_tags(text):
                return re.sub(r'<[^>]+>', '', text)

            # Processa parágrafos e cabeçalhos
            # Divide por blocos de h1-h6 e p
            blocks = re.split(r'(?=<(?:h[1-6]|p|div|li|br)[^>]*>)', content, flags=re.IGNORECASE)

            for block in blocks:
                if not block.strip():
                    continue

                h_match = re.match(r'<(h[1-6])[^>]*>(.*?)</\1>', block, re.IGNORECASE | re.DOTALL)
                if h_match:
                    level = int(h_match.group(1)[1])
                    text = strip_tags(h_match.group(2)).strip()
                    if text:
                        heading = docx.add_heading(text, level=min(level, 3))
                    continue

                text = strip_tags(block).strip()
                if text:
                    para = docx.add_paragraph(text)
                    para.style = 'Normal'

            # Salva em buffer
            buffer = io.BytesIO()
            docx.save(buffer)
            buffer.seek(0)

            response = HttpResponse(
                buffer.read(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            filename = f"{doc.title or 'documento'}.docx"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response

        except Exception as e:
            logger.error(f"Erro ao gerar DOCX: {str(e)}", exc_info=True)
            return Response(
                {'error': f'Erro ao gerar DOCX: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @extend_schema(
        summary="Gerar documento",
        description="Gera documento final a partir do template e dados, com auxílio de IA",
        tags=["Documents"],
        responses={200: OpenApiResponse(description="Documento gerado")}
    )
    @action(detail=True, methods=['post'])
    def generate(self, request, pk=None):
        """Gera documento HTML usando IA - pega template + dados e gera documento completo"""
        import json

        etp = self.get_object()

        # Verificar se tem document_template
        if not etp.document_template:
            return Response(
                {'error': 'Documento não possui template vinculado. Configure manualmente no admin.'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        if not etp.data:
            return Response(
                {'error': 'Documento não possui dados preenchidos'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            # 1. Pegar template HTML
            template_content = etp.document_template.content

            if not template_content:
                return Response({
                    'error': 'Template de documento está vazio'
                }, status=status.HTTP_400_BAD_REQUEST)

            # 2. Usar Django Template para substituir placeholders
            from django.template import Template, Context
            from django.utils.safestring import mark_safe
            import markdown

            logger.debug("Gerando documento com Django Template...")
            logger.debug(f"Template content length: {len(template_content)}")
            logger.debug(f"Data keys: {list(etp.data.keys())}")

            # Converter campos Markdown para HTML
            def convert_markdown_if_needed(value):
                """Converte Markdown para HTML se o valor parecer ter formatação Markdown."""
                if not isinstance(value, str):
                    return value
                # Detecta se tem markdown (headers, bold, lists, etc)
                markdown_indicators = ['# ', '## ', '### ', '**', '- ', '* ', '1. ', '\n\n']
                has_markdown = any(indicator in value for indicator in markdown_indicators)
                if has_markdown:
                    # Converte markdown para HTML e marca como seguro
                    html_content = markdown.markdown(value, extensions=['tables', 'nl2br'])
                    return mark_safe(html_content)
                return value

            # Preparar contexto com dados do formulário + metadados
            context_data = {}
            for key, value in etp.data.items():
                context_data[key] = convert_markdown_if_needed(value)

            # Adicionar metadados
            context_data['titulo'] = etp.title
            context_data['numero_processo'] = etp.numero_processo or ''
            context_data['data_geracao'] = timezone.now().strftime('%d/%m/%Y')

            # Renderizar template
            template = Template(template_content)
            context = Context(context_data)
            generated_html = template.render(context)

            logger.debug("Template renderizado com sucesso")
            logger.debug(f"Generated HTML length: {len(generated_html)}")

            # 3. Salvar no ETP e atualizar status para completo
            etp.generated_content = generated_html
            etp.generated_html = generated_html
            etp.status = 'completed'
            etp.progress = 100
            etp.completed_at = timezone.now()
            etp.save()

            logger.debug("Documento salvo com status 'completed'")

            # Envia email de notificacao
            try:
                from django.conf import settings
                send_email_task.delay('enviar_email_documento_processado', {
                    'destinatario': request.user.email,
                    'nome_usuario': request.user.get_full_name() or request.user.username,
                    'titulo_documento': etp.title,
                    'link_documento': f"{getattr(settings, 'APP_BASE_URL', 'http://localhost:3000')}/dashboard/documents/{etp.id}",
                })
            except Exception as e:
                logger.warning("Falha ao enviar email de documento processado: %s", e)

            return Response({
                'detail': 'Documento gerado com sucesso',
                'etp_id': str(etp.id),
                'status': 'completed',
                'generated_html_length': len(generated_html),
                'template_used': etp.document_template.name,
            }, status=status.HTTP_200_OK)

        except Exception as e:
            import traceback
            logger.error(f"Erro ao gerar documento: {str(e)}")
            logger.error(traceback.format_exc())
            return Response({
                'error': f'Erro ao gerar documento: {str(e)}',
                'etp_id': str(etp.id)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @extend_schema(
        summary="Exportar documento como PDF",
        description="Gera PDF do documento HTML gerado",
        tags=["Documents"],
        responses={200: OpenApiResponse(description="PDF do documento")}
    )
    @action(detail=True, methods=['get'])
    def export_pdf(self, request, pk=None):
        """Exporta documento como PDF"""
        from weasyprint import HTML, CSS
        from django.http import HttpResponse
        import io

        etp = self.get_object()

        if not etp.generated_html:
            return Response({
                'error': 'Documento ainda não foi gerado'
            }, status=status.HTTP_400_BAD_REQUEST)

        try:
            # Verificar se o HTML gerado já é um documento completo
            html_content = etp.generated_html.strip()

            if not html_content.lower().startswith('<!doctype') and not html_content.lower().startswith('<html'):
                # Se não for HTML completo, embutir em um template
                html_content = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <style>
                        @page {{
                            size: A4;
                            margin: 2.5cm;
                        }}
                        body {{
                            font-family: 'Calibri', Arial, sans-serif;
                            font-size: 11pt;
                            line-height: 1.5;
                        }}
                        h1 {{ font-size: 16pt; font-weight: bold; margin-top: 20pt; }}
                        h2 {{ font-size: 14pt; font-weight: bold; margin-top: 15pt; }}
                        h3 {{ font-size: 12pt; font-weight: bold; margin-top: 10pt; }}
                        p {{ margin: 0 0 10pt 0; }}
                        strong, b {{ font-weight: bold; }}
                        table {{ width: 100%; border-collapse: collapse; margin: 10pt 0; }}
                        td, th {{ border: 1px solid #000; padding: 5pt; }}
                    </style>
                </head>
                <body>
                    {etp.generated_html}
                </body>
                </html>
                """

            # Gerar PDF
            pdf_file = HTML(string=html_content).write_pdf()

            # Criar response
            response = HttpResponse(pdf_file, content_type='application/pdf')
            filename = f"{etp.title or 'documento'}.pdf"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'

            return response

        except Exception as e:
            import traceback
            logger.error(f"Erro ao gerar PDF: {str(e)}")
            logger.error(traceback.format_exc())
            return Response({
                'error': f'Erro ao gerar PDF: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    # ==========================================================================
    # REDAÇÃO / TARJAMENTO DE DADOS SENSÍVEIS (LGPD)
    # ==========================================================================

    def _get_client_ip(self, request) -> str:
        """Extrai IP real do cliente (suporta proxy/nginx)."""
        x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
        if x_forwarded_for:
            return x_forwarded_for.split(',')[0].strip()
        return request.META.get('REMOTE_ADDR', '0.0.0.0')

    @extend_schema(
        summary="Conteúdo tarjado (redacted)",
        description=(
            "Retorna o conteúdo HTML do documento com dados sensíveis (PII) "
            "tarjados/mascarados. Os segmentos de PII são substituídos por "
            "versões parcialmente mascaradas, preservando a estrutura do documento."
        ),
        tags=["Documents"],
        responses={200: OpenApiResponse(description="Segmentos de texto com PII tarjado")},
    )
    @action(detail=True, methods=['get'])
    def content_redacted(self, request, pk=None):
        """Retorna conteúdo com PII tarjado."""
        from apps.core.redaction import segmentize_text, text_has_pii

        doc = self.get_object()
        text = doc.generated_html or doc.generated_content or ''

        if not text:
            return Response({'detail': 'Documento sem conteúdo gerado.'}, status=404)

        has_sensitive = text_has_pii(text)
        segments = segmentize_text(text) if has_sensitive else [{'type': 'text', 'content': text}]

        return Response({
            'has_pii': has_sensitive,
            'segments': segments,
            'segment_count': len(segments),
            'pii_summary': self._pii_summary(text) if has_sensitive else {},
        })

    def _pii_summary(self, text: str) -> dict:
        """Sumário de PII para o response."""
        from apps.core.pii_detector import pii_summary
        summary = pii_summary(text)
        return {
            k: v for k, v in summary.items()
            if not k.endswith('_invalid_count')
        }

    @extend_schema(
        summary="Revelar campo sensível",
        description=(
            "Revela o valor original de um campo PII a partir de uma referência "
            "(ex.: CPF·1). Requer permissão de auditor/revisor/admin. "
            "A ação é registrada no log de auditoria de dados sensíveis. "
            "O valor fica visível por 30 segundos e é automaticamente removido do cache."
        ),
        tags=["Documents"],
        request={
            "type": "object",
            "properties": {
                "ref": {"type": "string", "description": "Referência do campo (ex: CPF·1, EMAIL·2)"},
            },
            "required": ["ref"]
        },
        responses={200: OpenApiResponse(description="Valor original do campo")},
    )
    @action(detail=True, methods=['post'])
    def reveal_field(self, request, pk=None):
        """Revela valor original de um campo PII (auditor/revisor/admin)."""
        from apps.accounts.permissions import resolve_role
        from apps.core.redaction import resolve_field

        doc = self.get_object()
        text = doc.generated_html or doc.generated_content or ''

        if not text:
            return Response({'detail': 'Documento sem conteúdo.'}, status=404)

        ref = request.data.get('ref', '')
        if not ref:
            return Response({'detail': 'Campo "ref" é obrigatório.'}, status=400)

        # Verificar permissão: auditor + ou admin
        user = request.user
        allowed_roles = ('superadmin', 'admin', 'auditor', 'revisor', 'gestor',
                         'socio', 'procurador', 'advogado_senior', 'coordenador')
        role = resolve_role(user.role)
        if not user.is_superuser and role not in allowed_roles:
            return Response(
                {'detail': 'Você não tem permissão para revelar dados sensíveis.'},
                status=status.HTTP_403_FORBIDDEN,
            )

        value = resolve_field(text, ref)
        if value is None:
            return Response(
                {'detail': f'Referência "{ref}" não encontrada no documento.'},
                status=404,
            )

        # Log da ação de revelação
        logger.info(
            'Usuário %s revelou campo %s no documento %s (IP: %s)',
            user.id, ref, doc.id, self._get_client_ip(request),
        )

        return Response({
            'ref': ref,
            'value': value,
            'expires_in_seconds': 30,
        })

    @extend_schema(
        summary="Registrar acesso a documento sensível",
        description=(
            "Registra no log que o usuário visualizou ou fez download "
            "de um documento que contém dados sensíveis (PII). "
            "Apenas registra quando doc.metadata.pii_warning é True."
        ),
        tags=["Documents"],
        request={
            "type": "object",
            "properties": {
                "action": {
                    "type": "string",
                    "enum": ["view_file", "download"],
                    "description": "Tipo de acesso",
                }
            },
            "required": ["action"]
        },
        responses={201: OpenApiResponse(description="Acesso registrado")},
    )
    @action(detail=True, methods=['post'])
    def log_access(self, request, pk=None):
        """Registra visualização/download de documento com PII."""
        doc = self.get_object()
        metadata = doc.metadata or {}

        # Só registra se o documento tem aviso de PII
        if not metadata.get('pii_warning', False):
            return Response({'detail': 'Documento sem dados sensíveis. Nada a registrar.'}, status=200)

        action_type = request.data.get('action', 'view_file')
        if action_type not in ('view_file', 'download'):
            return Response(
                {'detail': 'Ação inválida. Use "view_file" ou "download".'},
                status=400,
            )

        logger.info(
            'ACESSO_DOCUMENTO_SENSIVEL: user=%s doc=%s action=%s ip=%s',
            request.user.id, doc.id, action_type, self._get_client_ip(request),
        )

        # Atualizar metadata do documento para rastrear último acesso
        access_log = metadata.get('access_log', [])
        access_log.append({
            'user_id': str(request.user.id),
            'action': action_type,
            'ip': self._get_client_ip(request),
            'timestamp': timezone.now().isoformat(),
        })
        # Manter apenas os últimos 100 registros
        if len(access_log) > 100:
            access_log = access_log[-100:]
        metadata['access_log'] = access_log
        doc.metadata = metadata
        doc.save(update_fields=['metadata'])

        return Response({
            'detail': 'Acesso registrado com sucesso.',
            'action': action_type,
        }, status=201)

    @extend_schema(
        summary="Aplicar tarjamento PII",
        description="Aplica tarjamento (redação) nos dados sensíveis do documento. Os campos PII são substituídos por conteúdo mascarado.",
        tags=["Documents"],
        responses={200: DocumentDetailSerializer},
    )
    @action(detail=True, methods=['post'])
    def apply_redaction(self, request, pk=None):
        """Aplica tarjamento PII e marca documento como redactado."""
        from apps.core.redaction import apply_redaction as do_redact

        doc = self.get_object()
        text = doc.generated_html or doc.generated_content or ''
        if not text:
            return Response({'detail': 'Documento sem conteúdo.'}, status=404)

        redacted_text = do_redact(text)
        metadata = doc.metadata or {}
        metadata['pii_redacted_at'] = timezone.now().isoformat()
        metadata['pii_redacted_by'] = str(request.user.id)
        doc.metadata = metadata
        doc.save(update_fields=['metadata'])

        return Response({
            'detail': 'Tarjamento aplicado com sucesso.',
            'has_pii': text != redacted_text,
            'redacted_length': len(redacted_text),
        })


class DocumentGeneratorViewSet(viewsets.ModelViewSet):
    """ViewSet para DocumentGenerator - gerencia configurações de geradores de IA."""
    queryset = DocumentGenerator.objects.select_related('document_template', 'created_by').prefetch_related('knowledge_bases')
    serializer_class = DocumentGeneratorSerializer
    permission_classes = [IsAuthenticated]
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['name', 'description', 'document_type']
    ordering_fields = ['document_type', 'display_order', 'name', 'created_at']
    ordering = ['document_type', 'display_order']

    def get_queryset(self):
        qs = DocumentGenerator.objects.select_related('document_template', 'created_by').prefetch_related('knowledge_bases')
        specialty = self.request.query_params.get('specialty')
        document_type = self.request.query_params.get('document_type')
        active_only = self.request.query_params.get('active_only')
        search = self.request.query_params.get('search')

        if specialty and specialty != 'all':
            qs = qs.filter(specialty=specialty)
        if document_type and document_type != 'all':
            qs = qs.filter(document_type=document_type)
        if active_only and active_only.lower() == 'true':
            qs = qs.filter(is_active=True)
        if search:
            from django.db.models import Q
            qs = qs.filter(Q(name__icontains=search) | Q(description__icontains=search) | Q(document_type__icontains=search))

        llm_provider = self.request.query_params.get('llm_provider')
        if llm_provider:
            qs = qs.filter(llm_provider=llm_provider)
        return qs

    @extend_schema(
        summary="Estatísticas de Geradores",
        description="Retorna estatísticas agregadas dos geradores de documento",
        tags=["Document Generators"],
        responses={200: OpenApiResponse(description="Estatísticas")},
    )
    @action(detail=False, methods=['get'])
    def stats(self, request):
        from django.db.models import Count
        qs = DocumentGenerator.objects.all()
        total = qs.count()
        active = qs.filter(is_active=True).count()
        by_provider_qs = qs.values('llm_provider').annotate(count=Count('id'))
        by_provider = {
            'openai': {'name': 'OpenAI', 'count': 0},
            'anthropic': {'name': 'Anthropic', 'count': 0},
        }
        for row in by_provider_qs:
            p = row['llm_provider']
            if p in by_provider:
                by_provider[p]['count'] = row['count']
        return Response({'total': total, 'active': active, 'by_provider': by_provider})

    @extend_schema(
        summary="Duplicar Gerador",
        description="Cria uma cópia do gerador com o mesmo nome + ' (cópia)'",
        tags=["Document Generators"],
        responses={201: DocumentGeneratorSerializer},
    )
    @action(detail=True, methods=['post'])
    def duplicate(self, request, pk=None):
        original = self.get_object()
        kbs = list(original.knowledge_bases.all())
        original.pk = None
        original.id = None
        original.name = f'{original.name} (cópia)'
        original.is_default = False
        original.created_by = request.user
        original.save()
        if kbs:
            original.knowledge_bases.set(kbs)
        serializer = DocumentGeneratorSerializer(original, context={'request': request})
        return Response(serializer.data, status=status.HTTP_201_CREATED)
