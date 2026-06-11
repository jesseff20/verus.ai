"""
Document Export Service — Conversão de Markdown para DOCX e ODT.

Utiliza python-docx para DOCX e odfpy para ODT.
Aplica personalização do blueprint (fontes, cores, margens, capa).
"""
import html
import logging
import re
import markdown
from html.parser import HTMLParser
from io import BytesIO
from typing import Optional, Dict, Any, List, Tuple

from ..utils import normalize_subsection_breaks

logger = logging.getLogger(__name__)


# Tags HTML em bloco geradas pelo richtext (TinyMCE). Quando uma linha do
# conteudo do agente comeca com uma destas, bufferizamos e parseamos via
# _HTMLBlockParser ao inves de tratar como paragrafo de texto puro.
_HTML_BLOCK_OPEN_RE = re.compile(
    r'^<\s*(p|ul|ol|h[1-6]|blockquote|div)[\s>/]',
    re.IGNORECASE,
)


class DocumentExportService:
    """
    Converte Markdown em DOCX e ODT com estilização profissional
    baseada nas configurações do DocumentBlueprint.
    """

    # ─── HTML Table Parser ─────────────────────────────────────────

    @staticmethod
    def _parse_html_table(html: str) -> Optional[Dict[str, Any]]:
        """
        Parseia um bloco <table>...</table> e retorna um grid retangular que
        respeita rowspan/colspan, preserva todas as linhas (incluindo banners
        em <thead> e linhas vazias em <tbody>) e marca celulas absorvidas por
        spans como 'covered' — pra os renderers DOCX/ODT mesclarem nativamente.

        Formato de retorno:
          {
            'type': 'table',
            'cols': N,            # numero total de colunas (apos expandir colspans)
            'n_rows': M,          # numero total de linhas
            'head_count': K,      # numero de linhas iniciais que sao header
            'grid': [             # M x N — cada celula e None ou dict abaixo
                [
                    {
                        'is_origin': True/False,   # False = absorvida por span de outra
                        'is_header': True/False,
                        'colspan': int,            # so na origem
                        'rowspan': int,            # so na origem
                        'text': str,               # so na origem
                    },
                    ...
                ],
                ...
            ],
          }

        Decodifica entidades HTML (convert_charrefs=True) — &aacute; -> a etc.
        Linhas com apenas <th colspan=N> (banners) entram como header com a
        celula origem cobrindo as N colunas, em vez de serem descartadas.
        """
        class _TableParser(HTMLParser):
            def __init__(self):
                super().__init__(convert_charrefs=True)
                # Lista de (is_thead, [raw_cell, ...])
                self.raw_rows: List[Tuple[bool, List[Dict[str, Any]]]] = []
                self._in_thead = False
                self._in_tbody = False
                self._current_row: Optional[List[Dict[str, Any]]] = None
                self._cell_attrs: Optional[Dict[str, Any]] = None
                self._cell_buf: List[str] = []

            def handle_starttag(self, tag, attrs):
                tag = tag.lower()
                attrs_dict = dict(attrs)
                if tag == 'thead':
                    self._in_thead = True
                elif tag == 'tbody':
                    self._in_tbody = True
                    self._in_thead = False
                elif tag == 'tfoot':
                    self._in_thead = False
                elif tag == 'tr':
                    self._current_row = []
                elif tag in ('th', 'td'):
                    try:
                        cs = int(attrs_dict.get('colspan', '1') or '1')
                    except ValueError:
                        cs = 1
                    try:
                        rs = int(attrs_dict.get('rowspan', '1') or '1')
                    except ValueError:
                        rs = 1
                    self._cell_attrs = {
                        'is_th': tag == 'th',
                        'colspan': max(1, cs),
                        'rowspan': max(1, rs),
                    }
                    self._cell_buf = []
                elif tag == 'br' and self._cell_attrs is not None:
                    self._cell_buf.append('\n')

            def handle_endtag(self, tag):
                tag = tag.lower()
                if tag == 'thead':
                    self._in_thead = False
                elif tag == 'tbody':
                    self._in_tbody = False
                elif tag == 'tr':
                    if self._current_row is not None:
                        # Se nenhum thead/tbody explicito, primeiras linhas com
                        # apenas <th> sao tratadas como header (heuristica).
                        is_head = self._in_thead
                        if not self._in_thead and not self._in_tbody:
                            if all(c.get('is_th') for c in self._current_row) and self._current_row:
                                is_head = True
                        self.raw_rows.append((is_head, self._current_row))
                        self._current_row = None
                elif tag in ('th', 'td'):
                    if self._cell_attrs is None or self._current_row is None:
                        self._cell_attrs = None
                        return
                    text = ''.join(self._cell_buf)
                    # Normaliza espacos (incluindo NBSP) e mantem quebras
                    text = text.replace('\xa0', ' ')
                    # Junta espacos consecutivos preservando \n
                    text = re.sub(r'[ \t]+', ' ', text).strip()
                    self._current_row.append({
                        'text': text,
                        'is_th': self._cell_attrs['is_th'],
                        'colspan': self._cell_attrs['colspan'],
                        'rowspan': self._cell_attrs['rowspan'],
                    })
                    self._cell_attrs = None
                    self._cell_buf = []

            def handle_data(self, data):
                if self._cell_attrs is not None:
                    self._cell_buf.append(data)

        parser = _TableParser()
        try:
            parser.feed(html)
            parser.close()
        except Exception as exc:
            logger.warning("Falha ao parsear <table>: %s", exc)
            return None

        if not parser.raw_rows:
            return None

        # Determina numero de colunas: maior soma de colspans em uma linha,
        # considerando que rowspans de linhas anteriores ocupam posicoes.
        # Calculamos o grid de uma vez ja resolvendo as posicoes.
        n_rows = len(parser.raw_rows)

        # Primeira passada: estimar max_cols olhando cada linha + ocupacao.
        # Precisamos de um grid mutavel. Usamos lista de listas indexada
        # dinamicamente (cresce conforme necessario).
        grid: List[List[Optional[Dict[str, Any]]]] = [[] for _ in range(n_rows)]

        def _ensure_width(row_list, width):
            while len(row_list) < width:
                row_list.append(None)

        for r, (is_head, cells) in enumerate(parser.raw_rows):
            col = 0
            for cell in cells:
                # Pula posicoes ja ocupadas por rowspan de linha acima
                _ensure_width(grid[r], col + 1)
                while col < len(grid[r]) and grid[r][col] is not None:
                    col += 1
                _ensure_width(grid[r], col + 1)

                cs = cell['colspan']
                rs = min(cell['rowspan'], n_rows - r)

                origin_cell = {
                    'is_origin': True,
                    'is_header': cell['is_th'] or is_head,
                    'colspan': cs,
                    'rowspan': rs,
                    'text': cell['text'],
                }
                grid[r][col] = origin_cell

                # Marcar posicoes absorvidas — origin_row indica de qual
                # linha vem o span (importante pro render ODT distinguir
                # colspan-da-mesma-linha vs rowspan-de-linha-anterior).
                for dr in range(rs):
                    for dc in range(cs):
                        if dr == 0 and dc == 0:
                            continue
                        target_row = grid[r + dr]
                        _ensure_width(target_row, col + dc + 1)
                        if target_row[col + dc] is None:
                            target_row[col + dc] = {
                                'is_origin': False,
                                'origin_row': r,
                                'is_header': cell['is_th'] or is_head,
                            }
                col += cs

        # Normaliza largura de todas as linhas
        max_cols = max((len(row) for row in grid), default=0)
        for row in grid:
            _ensure_width(row, max_cols)

        head_count = 0
        for is_head, _ in parser.raw_rows:
            if is_head:
                head_count += 1
            else:
                break  # cabecalho vem antes de body

        if max_cols == 0:
            return None

        return {
            'type': 'table',
            'cols': max_cols,
            'n_rows': n_rows,
            'head_count': head_count,
            'grid': grid,
        }

    # ─── HTML em bloco (TinyMCE/richtext) → blocks estruturados ─────

    @staticmethod
    def _parse_html_blocks(html_text: str) -> List[Dict[str, Any]]:
        """
        Parseia HTML em bloco (gerado por TinyMCE) em multiplos blocks
        estruturados (paragraph, list_item, heading), preservando
        formatacao inline (bold/italic) e decodificando entidades HTML
        (&aacute;, &nbsp;, etc.) automaticamente.

        Tabelas continuam sendo tratadas pelo caminho dedicado
        `_parse_html_table` — se um <table> aparece embutido aqui, e
        coletado e parseado em separado pra preservar headers/rows.
        """
        class _HTMLBlockParser(HTMLParser):
            def __init__(self):
                # convert_charrefs=True faz HTMLParser decodificar entidades
                # (&aacute; -> á, &nbsp; -> espaco) antes de chamar handle_data.
                super().__init__(convert_charrefs=True)
                self.blocks: List[Dict[str, Any]] = []
                self._current_type: Optional[str] = None
                self._current_runs: List[Dict[str, Any]] = []
                self._text_buf: List[str] = []
                self._heading_level = 1
                self._ordered = False
                self._bold_depth = 0
                self._italic_depth = 0
                # Buffer pra capturar HTML cru de <table>...</table> e parsear
                # com _parse_html_table (preserva headers/rows).
                self._table_buf: Optional[List[str]] = None
                self._table_depth = 0

            def _flush_run(self):
                if not self._text_buf:
                    return
                text = ''.join(self._text_buf)
                if not text:
                    self._text_buf = []
                    return
                self._current_runs.append({
                    'text': text,
                    'bold': self._bold_depth > 0,
                    'italic': self._italic_depth > 0,
                })
                self._text_buf = []

            def _close_block(self):
                self._flush_run()
                if self._current_type and self._current_runs:
                    full_text = ''.join(r['text'] for r in self._current_runs).strip()
                    if full_text:
                        block: Dict[str, Any] = {
                            'type': self._current_type,
                            'text': full_text,
                            'runs': self._current_runs,
                        }
                        if self._current_type == 'heading':
                            block['level'] = self._heading_level
                        elif self._current_type == 'list_item':
                            block['ordered'] = self._ordered
                        self.blocks.append(block)
                self._current_type = None
                self._current_runs = []
                self._text_buf = []

            # ── Captura de <table> embutida ──
            def handle_starttag(self, tag, attrs):
                tag = tag.lower()
                if self._table_buf is not None:
                    # Reconstrui o HTML da tabela pra repassar ao table parser.
                    attr_str = ''.join(
                        f' {k}="{v}"' for k, v in attrs if v is not None
                    )
                    self._table_buf.append(f'<{tag}{attr_str}>')
                    if tag == 'table':
                        self._table_depth += 1
                    return
                if tag == 'table':
                    self._close_block()
                    self._table_buf = ['<table>']
                    self._table_depth = 1
                    return

                if tag == 'p':
                    self._close_block()
                    self._current_type = 'paragraph'
                elif tag in ('ul', 'ol'):
                    self._close_block()
                    self._ordered = (tag == 'ol')
                elif tag == 'li':
                    self._close_block()
                    self._current_type = 'list_item'
                elif tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                    self._close_block()
                    self._current_type = 'heading'
                    self._heading_level = int(tag[1])
                elif tag == 'blockquote':
                    self._close_block()
                    self._current_type = 'paragraph'
                elif tag in ('strong', 'b'):
                    self._flush_run()
                    self._bold_depth += 1
                elif tag in ('em', 'i'):
                    self._flush_run()
                    self._italic_depth += 1
                elif tag == 'br':
                    if self._current_type is None:
                        self._current_type = 'paragraph'
                    self._text_buf.append('\n')
                # demais inline (span, a, code) ignoradas — texto continua

            def handle_endtag(self, tag):
                tag = tag.lower()
                if self._table_buf is not None:
                    self._table_buf.append(f'</{tag}>')
                    if tag == 'table':
                        self._table_depth -= 1
                        if self._table_depth <= 0:
                            html_table = ''.join(self._table_buf)
                            parsed = DocumentExportService._parse_html_table(html_table)
                            if parsed:
                                self.blocks.append(parsed)
                            self._table_buf = None
                            self._table_depth = 0
                    return

                if tag == 'p' or tag == 'li' or tag == 'blockquote':
                    self._close_block()
                elif tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                    self._close_block()
                elif tag in ('strong', 'b'):
                    self._flush_run()
                    self._bold_depth = max(0, self._bold_depth - 1)
                elif tag in ('em', 'i'):
                    self._flush_run()
                    self._italic_depth = max(0, self._italic_depth - 1)

            def handle_data(self, data):
                if self._table_buf is not None:
                    # Escapa de volta pra preservar fidelidade no re-parse.
                    self._table_buf.append(html.escape(data, quote=False))
                    return
                if self._current_type is None:
                    if not data.strip():
                        return
                    self._current_type = 'paragraph'
                self._text_buf.append(data)

            def close(self):
                super().close()
                self._close_block()

        parser = _HTMLBlockParser()
        try:
            parser.feed(html_text)
            parser.close()
        except Exception as exc:
            logger.warning("Falha ao parsear HTML em bloco: %s", exc)
            return []
        return parser.blocks

    # ─── Markdown → tokens intermediários ───────────────────────────

    @staticmethod
    def _parse_markdown_blocks(md_text: str) -> List[Dict[str, Any]]:
        """
        Parseia Markdown em blocos estruturados para alimentar
        os geradores DOCX/ODT sem depender de HTML intermediário.

        Retorna lista de dicts:
          {'type': 'heading', 'level': 2, 'text': '...'}
          {'type': 'paragraph', 'runs': [{'text':'...', 'bold':bool, 'italic':bool}]}
          {'type': 'list_item', 'ordered': bool, 'text': '...', 'runs': [...]}
          {'type': 'table', 'headers': [...], 'rows': [[...]]}
        """
        blocks: List[Dict[str, Any]] = []
        lines = md_text.split('\n')
        i = 0
        table_buf: List[str] = []

        def _flush_table():
            """Markdown pipe table → grid retangular trivial (sem spans)."""
            nonlocal table_buf
            if not table_buf:
                return
            headers = [c.strip() for c in table_buf[0].strip('|').split('|')]
            data_rows = []
            for row_line in table_buf[2:]:  # skip header + separator
                cells = [c.strip() for c in row_line.strip('|').split('|')]
                data_rows.append(cells)

            cols = max(len(headers), max((len(r) for r in data_rows), default=0))
            grid: List[List[Optional[Dict[str, Any]]]] = []
            # Header row
            head_row: List[Optional[Dict[str, Any]]] = []
            for c in range(cols):
                text = headers[c] if c < len(headers) else ''
                head_row.append({
                    'is_origin': True, 'is_header': True,
                    'colspan': 1, 'rowspan': 1, 'text': text,
                })
            grid.append(head_row)
            # Data rows
            for r in data_rows:
                row: List[Optional[Dict[str, Any]]] = []
                for c in range(cols):
                    text = r[c] if c < len(r) else ''
                    row.append({
                        'is_origin': True, 'is_header': False,
                        'colspan': 1, 'rowspan': 1, 'text': text,
                    })
                grid.append(row)

            blocks.append({
                'type': 'table',
                'cols': cols,
                'n_rows': len(grid),
                'head_count': 1,
                'grid': grid,
            })
            table_buf = []

        html_table_buf: List[str] = []
        in_html_table = False

        def _flush_html_table():
            nonlocal html_table_buf, in_html_table
            if not html_table_buf:
                return
            html_str = '\n'.join(html_table_buf)
            parsed = DocumentExportService._parse_html_table(html_str)
            if parsed:
                blocks.append(parsed)
            html_table_buf = []
            in_html_table = False

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # HTML table detection (caminho dedicado, preserva headers/rows)
            if not in_html_table and re.match(r'^<table[\s>]', stripped, re.IGNORECASE):
                _flush_table()
                in_html_table = True
                html_table_buf = [stripped]
                if re.search(r'</table>', stripped, re.IGNORECASE):
                    _flush_html_table()
                i += 1
                continue
            if in_html_table:
                html_table_buf.append(stripped)
                if re.search(r'</table>', stripped, re.IGNORECASE):
                    _flush_html_table()
                i += 1
                continue

            # HTML em bloco (TinyMCE/richtext): <p>, <ul>, <ol>, <h1-6>, etc.
            # Bufferiza ate linha em branco ou inicio de markdown puro, depois
            # parseia tudo de uma vez via _parse_html_blocks (entidades HTML
            # sao decodificadas dentro do parser).
            if _HTML_BLOCK_OPEN_RE.match(stripped):
                _flush_table()
                html_buf = [stripped]
                i += 1
                while i < len(lines):
                    nxt = lines[i].strip()
                    if not nxt:
                        i += 1
                        break
                    if nxt.startswith('#') or re.match(r'^\d+[.)]\s', nxt):
                        # Voltou a ser markdown puro — interrompe o bloco HTML.
                        break
                    html_buf.append(nxt)
                    i += 1
                html_full = '\n'.join(html_buf)
                blocks.extend(DocumentExportService._parse_html_blocks(html_full))
                continue

            # Markdown pipe table detection
            if stripped.startswith('|') and '|' in stripped[1:]:
                table_buf.append(stripped)
                i += 1
                continue
            else:
                _flush_table()

            # Empty line
            if not stripped:
                i += 1
                continue

            # Headings
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                blocks.append({'type': 'heading', 'level': level, 'text': text})
                i += 1
                continue

            # Unordered list
            list_match = re.match(r'^[-*+]\s+(.+)$', stripped)
            if list_match:
                text = list_match.group(1)
                blocks.append({
                    'type': 'list_item',
                    'ordered': False,
                    'text': text,
                    'runs': DocumentExportService._parse_inline(text),
                })
                i += 1
                continue

            # Ordered list
            olist_match = re.match(r'^\d+[.)]\s+(.+)$', stripped)
            if olist_match:
                text = olist_match.group(1)
                blocks.append({
                    'type': 'list_item',
                    'ordered': True,
                    'text': text,
                    'runs': DocumentExportService._parse_inline(text),
                })
                i += 1
                continue

            # Horizontal rule
            if re.match(r'^[-*_]{3,}\s*$', stripped):
                blocks.append({'type': 'hr'})
                i += 1
                continue

            # Regular paragraph (may span multiple lines)
            para_lines = [stripped]
            i += 1
            while i < len(lines):
                nxt = lines[i].strip()
                if not nxt or nxt.startswith('#') or nxt.startswith('|') or re.match(r'^[-*+]\s', nxt) or re.match(r'^\d+[.)]\s', nxt) or re.match(r'^[-*_]{3,}$', nxt):
                    break
                para_lines.append(nxt)
                i += 1
            full_text = ' '.join(para_lines)
            blocks.append({
                'type': 'paragraph',
                'text': full_text,
                'runs': DocumentExportService._parse_inline(full_text),
            })

        _flush_table()
        _flush_html_table()
        return blocks

    @staticmethod
    def _parse_inline(text: str) -> List[Dict[str, Any]]:
        """
        Parseia formatação inline (bold, italic, bold+italic) em runs.
        """
        runs: List[Dict[str, Any]] = []
        # Pattern: ***bold+italic***, **bold**, *italic*, normal text
        pattern = re.compile(
            r'(\*\*\*(.+?)\*\*\*)'   # bold+italic
            r'|(\*\*(.+?)\*\*)'       # bold
            r'|(\*(.+?)\*)'           # italic
            r'|(__(.+?)__)'            # bold (underscore)
            r'|(_(.+?)_)'             # italic (underscore)
            r'|([^*_]+)'              # plain text
        )
        for m in pattern.finditer(text):
            if m.group(2):  # bold+italic ***
                runs.append({'text': m.group(2), 'bold': True, 'italic': True})
            elif m.group(4):  # bold **
                runs.append({'text': m.group(4), 'bold': True, 'italic': False})
            elif m.group(6):  # italic *
                runs.append({'text': m.group(6), 'bold': False, 'italic': True})
            elif m.group(8):  # bold __
                runs.append({'text': m.group(8), 'bold': True, 'italic': False})
            elif m.group(10):  # italic _
                runs.append({'text': m.group(10), 'bold': False, 'italic': True})
            elif m.group(11):  # plain
                runs.append({'text': m.group(11), 'bold': False, 'italic': False})
        return runs if runs else [{'text': text, 'bold': False, 'italic': False}]

    # ─── DOCX Generation ───────────────────────────────────────────

    def generate_docx(
        self,
        markdown_content: str,
        objective: str = '',
        blueprint_config: Optional[Dict[str, Any]] = None,
    ) -> Optional[bytes]:
        """
        Converte Markdown para DOCX com estilização do blueprint.

        Returns:
            Bytes do arquivo .docx ou None se falhar.
        """
        markdown_content = normalize_subsection_breaks(markdown_content)
        try:
            from docx import Document
            from docx.shared import Pt, Cm, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT

            cfg = blueprint_config or {}
            doc = Document()

            # ── Page setup ──
            section = doc.sections[0]
            section.page_width = Cm(21)    # A4
            section.page_height = Cm(29.7)
            section.top_margin = self._parse_cm(cfg.get('pdf_page_margin_top', '2.5cm'))
            section.bottom_margin = self._parse_cm(cfg.get('pdf_page_margin_bottom', '2.5cm'))
            section.left_margin = self._parse_cm(cfg.get('pdf_page_margin_left', '3cm'))
            section.right_margin = self._parse_cm(cfg.get('pdf_page_margin_right', '2cm'))

            # ── Typography config ──
            font_name = cfg.get('pdf_font_family', 'Times New Roman')
            font_size = self._parse_pt(cfg.get('pdf_font_size', '12pt'))
            primary_color = self._hex_to_rgb(cfg.get('primary_color', '#7030A0'))
            secondary_color = self._hex_to_rgb(cfg.get('secondary_color', '#5B2EE0'))

            # ── Default style ──
            style = doc.styles['Normal']
            style.font.name = font_name
            style.font.size = Pt(font_size)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # ── Cover page ──
            if cfg.get('cover_page_enabled', False):
                self._add_docx_cover(doc, objective, cfg, primary_color, font_name)

            # ── Document header (when no cover) ──
            if not cfg.get('cover_page_enabled', False):
                header_text = cfg.get('header_text', '')
                if header_text:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(header_text)
                    run.bold = True
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.name = font_name

                doc.add_paragraph()  # spacer

            # ── Content ──
            blocks = self._parse_markdown_blocks(markdown_content)
            # Document content always uses black text (judiciary standard)
            black_rgb = (0, 0, 0)
            self._render_blocks_docx(doc, blocks, font_name, font_size, black_rgb, black_rgb)

            # ── Footer text ──
            footer_text = cfg.get('footer_text', '')
            if footer_text:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(footer_text)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(102, 102, 102)
                run.font.name = font_name

            # ── Serialize ──
            buf = BytesIO()
            doc.save(buf)
            result = buf.getvalue()
            logger.info(f"DOCX gerado: {len(result)} bytes")
            return result

        except Exception as e:
            logger.error(f"Erro ao gerar DOCX: {e}", exc_info=True)
            return None

    def _add_docx_cover(self, doc, objective, cfg, primary_color, font_name):
        """Adiciona página de rosto ao DOCX."""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Organization
        org_text = cfg.get('cover_organization_text', cfg.get('organization_name', ''))
        if org_text:
            for _ in range(4):
                doc.add_paragraph()  # spacing top
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(org_text.upper())
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = font_name

        # Title
        doc.add_paragraph()
        title = cfg.get('cover_title', cfg.get('header_text', 'DOCUMENTO'))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = font_name

        # Subtitle
        subtitle = cfg.get('cover_subtitle', '')
        if subtitle:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(subtitle)
            run.italic = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(102, 102, 102)
            run.font.name = font_name

        # Footer
        cover_footer = cfg.get('cover_footer_text', '')
        if cover_footer:
            for _ in range(6):
                doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cover_footer)
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(102, 102, 102)
            run.font.name = font_name

        # Page break after cover
        doc.add_page_break()

    def _render_blocks_docx(self, doc, blocks, font_name, font_size, primary_color, secondary_color):
        """Renderiza blocos parseados em elementos DOCX."""
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        for block in blocks:
            btype = block['type']

            if btype == 'heading':
                level = min(block['level'], 4)
                p = doc.add_heading(block['text'], level=level)
                for run in p.runs:
                    run.font.name = font_name
                    run.font.color.rgb = RGBColor(0, 0, 0)

            elif btype == 'paragraph':
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1.5)
                for run_data in block.get('runs', []):
                    run = p.add_run(run_data['text'])
                    run.bold = run_data.get('bold', False)
                    run.italic = run_data.get('italic', False)
                    run.font.name = font_name
                    run.font.size = Pt(font_size)

            elif btype == 'list_item':
                style = 'List Number' if block.get('ordered') else 'List Bullet'
                p = doc.add_paragraph(style=style)
                for run_data in block.get('runs', []):
                    run = p.add_run(run_data['text'])
                    run.bold = run_data.get('bold', False)
                    run.italic = run_data.get('italic', False)
                    run.font.name = font_name
                    run.font.size = Pt(font_size)

            elif btype == 'table':
                self._render_table_docx_grid(doc, block, font_name)

            elif btype == 'hr':
                doc.add_paragraph()  # visual separator

    def _render_table_docx_grid(self, doc, block, font_name):
        """
        Renderiza tabela a partir do grid retangular (com spans).
        Cria tabela completa N x M, popula celulas-origem e mescla regioes
        com python-docx merge — preserva banners (colspan), headers
        compostos (rowspan/colspan) e linhas vazias do template.
        """
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        cols = block.get('cols', 0)
        n_rows = block.get('n_rows', 0)
        grid = block.get('grid')
        if not cols or not n_rows or not grid:
            return

        table = doc.add_table(rows=n_rows, cols=cols)
        table.style = 'Table Grid'
        table.autofit = True

        for r in range(n_rows):
            for c in range(cols):
                cell_data = grid[r][c] if c < len(grid[r]) else None
                if cell_data is None or not cell_data.get('is_origin'):
                    continue

                cs = cell_data.get('colspan', 1)
                rs = cell_data.get('rowspan', 1)
                is_header = cell_data.get('is_header', False)
                text = cell_data.get('text', '') or ''

                origin = table.cell(r, c)
                # Mesclar antes de setar texto evita perda em algumas versoes
                if cs > 1 or rs > 1:
                    end_r = min(r + rs - 1, n_rows - 1)
                    end_c = min(c + cs - 1, cols - 1)
                    end = table.cell(end_r, end_c)
                    origin = origin.merge(end)

                # Limpa qualquer paragrafo default e define texto
                origin.text = text

                for p in origin.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.font.name = font_name
                        run.font.size = Pt(10)
                        if is_header:
                            run.bold = True

                if is_header:
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'd3d3d3')
                    shading.set(qn('w:val'), 'clear')
                    origin._tc.get_or_add_tcPr().append(shading)

    # ─── ODT Generation ────────────────────────────────────────────

    def generate_odt(
        self,
        markdown_content: str,
        objective: str = '',
        blueprint_config: Optional[Dict[str, Any]] = None,
    ) -> Optional[bytes]:
        """
        Converte Markdown para ODT com estilização do blueprint.

        Returns:
            Bytes do arquivo .odt ou None se falhar.
        """
        markdown_content = normalize_subsection_breaks(markdown_content)
        try:
            from odf.opendocument import OpenDocumentText
            from odf.style import (
                Style, TextProperties, ParagraphProperties,
                TableProperties, TableColumnProperties, TableCellProperties,
                PageLayoutProperties, MasterPage, PageLayout,
            )
            from odf.text import P, H, List as OdfList, ListItem, Span
            from odf.table import Table, TableColumn, TableRow, TableCell

            cfg = blueprint_config or {}
            odt = OpenDocumentText()

            # ── Typography config ──
            font_name = cfg.get('pdf_font_family', 'Times New Roman')
            font_size_str = cfg.get('pdf_font_size', '12pt')
            primary_hex = cfg.get('primary_color', '#7030A0')
            secondary_hex = cfg.get('secondary_color', '#5B2EE0')

            # ── Page layout ──
            pl = PageLayout(name='PageLayout')
            pl.addElement(PageLayoutProperties(
                pagewidth='21cm',
                pageheight='29.7cm',
                margintop=cfg.get('pdf_page_margin_top', '2.5cm'),
                marginbottom=cfg.get('pdf_page_margin_bottom', '2.5cm'),
                marginleft=cfg.get('pdf_page_margin_left', '3cm'),
                marginright=cfg.get('pdf_page_margin_right', '2cm'),
            ))
            odt.automaticstyles.addElement(pl)

            mp = MasterPage(name='Standard', pagelayoutname=pl)
            odt.masterstyles.addElement(mp)

            # ── Styles ──
            # Normal paragraph
            normal_style = Style(name='NormalParagraph', family='paragraph')
            normal_style.addElement(ParagraphProperties(
                textalign='justify',
                textindent='1.5cm',
                marginbottom='0.4cm',
            ))
            normal_style.addElement(TextProperties(
                fontname=font_name,
                fontsize=font_size_str,
                color='#000000',
            ))
            odt.styles.addElement(normal_style)

            # Heading styles
            for level in range(1, 5):
                sizes = {'1': '16pt', '2': '14pt', '3': '12pt', '4': '11pt'}
                hs = Style(name=f'Heading{level}', family='paragraph')
                hs.addElement(ParagraphProperties(
                    textalign='left' if level > 1 else 'center',
                    marginbottom='0.3cm',
                    margintop='0.6cm',
                    keepwithnext='always',
                ))
                hs.addElement(TextProperties(
                    fontname=font_name,
                    fontsize=sizes.get(str(level), '12pt'),
                    fontweight='bold',
                    color='#000000',
                ))
                odt.styles.addElement(hs)

            # Bold span style
            bold_style = Style(name='BoldSpan', family='text')
            bold_style.addElement(TextProperties(fontweight='bold'))
            odt.styles.addElement(bold_style)

            # Italic span style
            italic_style = Style(name='ItalicSpan', family='text')
            italic_style.addElement(TextProperties(fontstyle='italic'))
            odt.styles.addElement(italic_style)

            # Bold+italic span
            bold_italic_style = Style(name='BoldItalicSpan', family='text')
            bold_italic_style.addElement(TextProperties(fontweight='bold', fontstyle='italic'))
            odt.styles.addElement(bold_italic_style)

            # Center style (for cover/header)
            center_style = Style(name='CenterParagraph', family='paragraph')
            center_style.addElement(ParagraphProperties(textalign='center'))
            center_style.addElement(TextProperties(fontname=font_name, fontsize=font_size_str))
            odt.styles.addElement(center_style)

            # Table anchor style: paragrafo OBRIGATORIO depois de toda tabela
            # (regra do ODT Writer — nao da pra remover). Usar font 2pt + zero
            # margens deixa o paragrafo praticamente invisivel (~0.7mm) e evita
            # que OpenOffice/Writer renderize uma "linha extra" notavel ao final
            # de cada tabela. Word/Office 365 ja renderizava sem problemas.
            anchor_style = Style(name='TableAnchor', family='paragraph')
            anchor_style.addElement(ParagraphProperties(
                margintop='0cm',
                marginbottom='0cm',
                textindent='0cm',
            ))
            anchor_style.addElement(TextProperties(fontsize='2pt'))
            odt.styles.addElement(anchor_style)

            # Paragrafo dentro de celula de header (centralizado)
            cell_header_para = Style(name='TableHeaderParagraph', family='paragraph')
            cell_header_para.addElement(ParagraphProperties(
                textalign='center', margintop='0cm', marginbottom='0cm', textindent='0cm',
            ))
            cell_header_para.addElement(TextProperties(fontname=font_name, fontsize='10pt'))
            odt.automaticstyles.addElement(cell_header_para)

            # Paragrafo dentro de celula normal (esquerda, sem indent)
            cell_para = Style(name='TableCellParagraph', family='paragraph')
            cell_para.addElement(ParagraphProperties(
                textalign='left', margintop='0cm', marginbottom='0cm', textindent='0cm',
            ))
            cell_para.addElement(TextProperties(fontname=font_name, fontsize='10pt'))
            odt.automaticstyles.addElement(cell_para)

            # Table cell style
            tc_style = Style(name='TableCell', family='table-cell')
            tc_style.addElement(TableCellProperties(padding='0.2cm', border='0.05pt solid #cccccc'))
            odt.automaticstyles.addElement(tc_style)

            # Table header cell style
            th_style = Style(name='TableHeaderCell', family='table-cell')
            th_style.addElement(TableCellProperties(
                padding='0.2cm',
                border='0.05pt solid #cccccc',
                backgroundcolor='#d3d3d3',
            ))
            odt.automaticstyles.addElement(th_style)

            # Table style
            table_style = Style(name='DocTable', family='table')
            table_style.addElement(TableProperties(width='17cm', align='margins'))
            odt.automaticstyles.addElement(table_style)

            # ── Cover page ──
            if cfg.get('cover_page_enabled', False):
                self._add_odt_cover(odt, objective, cfg, primary_hex, font_name)

            # ── Document header (when no cover) ──
            if not cfg.get('cover_page_enabled', False):
                header_text = cfg.get('header_text', '')
                if header_text:
                    h = H(outlinelevel=1, stylename='Heading1')
                    h.addText(header_text.upper())
                    odt.text.addElement(h)

                odt.text.addElement(P())  # spacer

            # ── Content ──
            blocks = self._parse_markdown_blocks(markdown_content)
            self._render_blocks_odt(odt, blocks)

            # ── Footer ──
            footer_text = cfg.get('footer_text', '')
            if footer_text:
                p = P(stylename='CenterParagraph')
                p.addText(footer_text)
                odt.text.addElement(p)

            # Paragrafo OBRIGATORIO no fim do body — regra do ODT: documento
            # nao pode terminar com uma tabela. Sempre adiciona um paragrafo
            # NormalParagraph (12pt) no final, com conteudo invisivel mas
            # presente. Sem ele, OpenOffice/Writer renderiza um retangulo
            # fantasma "absorvendo" a tabela.
            trailing = P(stylename='NormalParagraph')
            trailing.addText(' ')  # caractere invisivel — forca renderizacao real
            odt.text.addElement(trailing)

            # ── Serialize ──
            buf = BytesIO()
            odt.save(buf)
            result = buf.getvalue()
            logger.info(f"ODT gerado: {len(result)} bytes")
            return result

        except Exception as e:
            logger.error(f"Erro ao gerar ODT: {e}", exc_info=True)
            return None

    def _add_odt_cover(self, odt, objective, cfg, primary_hex, font_name):
        """Adiciona página de rosto ao ODT."""
        from odf.text import P, H
        from odf.style import Style, ParagraphProperties, TextProperties

        # Cover title style
        cover_title_style = Style(name='CoverTitle', family='paragraph')
        cover_title_style.addElement(ParagraphProperties(
            textalign='center',
            margintop='6cm',
            marginbottom='1cm',
        ))
        cover_title_style.addElement(TextProperties(
            fontname=font_name,
            fontsize='24pt',
            fontweight='bold',
            color='#000000',
        ))
        odt.styles.addElement(cover_title_style)

        # Cover org style
        cover_org_style = Style(name='CoverOrg', family='paragraph')
        cover_org_style.addElement(ParagraphProperties(textalign='center', marginbottom='2cm'))
        cover_org_style.addElement(TextProperties(
            fontname=font_name,
            fontsize='14pt',
            fontweight='bold',
            color='#000000',
        ))
        odt.styles.addElement(cover_org_style)

        # Page break style
        pb_style = Style(name='PageBreakParagraph', family='paragraph')
        pb_style.addElement(ParagraphProperties(breakafter='page'))
        odt.styles.addElement(pb_style)

        # Organization
        org_text = cfg.get('cover_organization_text', cfg.get('organization_name', ''))
        if org_text:
            p = P(stylename='CoverOrg')
            p.addText(org_text.upper())
            odt.text.addElement(p)

        # Title
        title = cfg.get('cover_title', cfg.get('header_text', 'DOCUMENTO'))
        p = P(stylename='CoverTitle')
        p.addText(title.upper())
        odt.text.addElement(p)

        # Subtitle
        subtitle = cfg.get('cover_subtitle', '')
        if subtitle:
            p = P(stylename='CenterParagraph')
            p.addText(subtitle)
            odt.text.addElement(p)

        # Page break
        p = P(stylename='PageBreakParagraph')
        odt.text.addElement(p)

    def _render_blocks_odt(self, odt, blocks):
        """Renderiza blocos parseados em elementos ODT."""
        from odf.text import P, H, Span, List as OdfList, ListItem
        from odf.table import (
            Table, TableColumn, TableRow, TableCell, CoveredTableCell,
        )

        for block in blocks:
            btype = block['type']

            if btype == 'heading':
                level = min(block['level'], 4)
                h = H(outlinelevel=level, stylename=f'Heading{level}')
                h.addText(block['text'])
                odt.text.addElement(h)

            elif btype == 'paragraph':
                p = P(stylename='NormalParagraph')
                for run_data in block.get('runs', []):
                    if run_data.get('bold') and run_data.get('italic'):
                        s = Span(stylename='BoldItalicSpan')
                        s.addText(run_data['text'])
                        p.addElement(s)
                    elif run_data.get('bold'):
                        s = Span(stylename='BoldSpan')
                        s.addText(run_data['text'])
                        p.addElement(s)
                    elif run_data.get('italic'):
                        s = Span(stylename='ItalicSpan')
                        s.addText(run_data['text'])
                        p.addElement(s)
                    else:
                        p.addText(run_data['text'])
                odt.text.addElement(p)

            elif btype == 'list_item':
                p = P(stylename='NormalParagraph')
                bullet = '• ' if not block.get('ordered') else ''
                for run_data in block.get('runs', []):
                    text = run_data['text']
                    if run_data.get('bold'):
                        s = Span(stylename='BoldSpan')
                        s.addText(bullet + text if bullet else text)
                        p.addElement(s)
                        bullet = ''
                    elif run_data.get('italic'):
                        s = Span(stylename='ItalicSpan')
                        s.addText(bullet + text if bullet else text)
                        p.addElement(s)
                        bullet = ''
                    else:
                        p.addText(bullet + text)
                        bullet = ''
                odt.text.addElement(p)

            elif btype == 'table':
                self._render_table_odt_grid(odt, block)

            elif btype == 'hr':
                odt.text.addElement(P())  # visual separator

    def _render_table_odt_grid(self, odt, block):
        """
        Renderiza tabela ODT respeitando rowspan/colspan (mesmo comportamento
        do DOCX). Estrategia per ODF spec, forma 1:
          - Cell origem com colspan/rowspan: emite numbercolumnsspanned /
            numberrowsspanned no <table:table-cell>.
          - Posicao absorvida por colspan da MESMA linha: pula (a cell
            origem ja cobre a posicao implicitamente).
          - Posicao absorvida por rowspan de linha ANTERIOR: emite
            <table:covered-table-cell/> (necessario pra manter o numero
            de elementos consistente nas linhas seguintes).

        Paragrafo obrigatorio depois da tabela (regra do formato ODT,
        forum.openoffice.org/viewtopic.php?t=2994) e adicionado UMA UNICA
        VEZ no fim do body em generate_odt — nao aqui.
        """
        from odf.text import P, Span
        from odf.table import (
            Table, TableColumn, TableRow, TableCell, CoveredTableCell,
        )

        cols = block.get('cols', 0)
        n_rows = block.get('n_rows', 0)
        grid = block.get('grid')
        if not cols or not n_rows or not grid:
            return

        table = Table(stylename='DocTable')
        for _ in range(cols):
            table.addElement(TableColumn())

        for r in range(n_rows):
            tr = TableRow()
            for c in range(cols):
                cell_data = grid[r][c] if c < len(grid[r]) else None

                if cell_data is None:
                    # Posicao sem dado — emite cell vazia regular.
                    tc = TableCell(stylename='TableCell')
                    tc.addElement(P(stylename='TableCellParagraph'))
                    tr.addElement(tc)
                    continue

                if not cell_data.get('is_origin'):
                    # Absorvida por algum span. Distingue origem:
                    #  - mesma linha (colspan): pula (forma 1 do spec)
                    #  - linha anterior (rowspan): emite covered-table-cell
                    if cell_data.get('origin_row', r) < r:
                        tr.addElement(CoveredTableCell())
                    continue

                # Cell origem
                cs = cell_data.get('colspan', 1)
                rs = cell_data.get('rowspan', 1)
                is_header = cell_data.get('is_header', False)
                text = cell_data.get('text', '') or ''
                style = 'TableHeaderCell' if is_header else 'TableCell'
                para_style = 'TableHeaderParagraph' if is_header else 'TableCellParagraph'

                kwargs = {'stylename': style}
                if cs > 1:
                    kwargs['numbercolumnsspanned'] = cs
                if rs > 1:
                    kwargs['numberrowsspanned'] = rs

                tc = TableCell(**kwargs)
                p = P(stylename=para_style)
                if text:
                    if is_header:
                        s = Span(stylename='BoldSpan')
                        s.addText(text)
                        p.addElement(s)
                    else:
                        p.addText(text)
                tc.addElement(p)
                tr.addElement(tc)
            table.addElement(tr)

        odt.text.addElement(table)

    # ─── Helpers ───────────────────────────────────────────────────

    @staticmethod
    def _parse_cm(value: str):
        """Converte '2.5cm' para docx.shared.Cm(2.5)."""
        from docx.shared import Cm
        num = float(re.sub(r'[^0-9.]', '', value) or '2.5')
        return Cm(num)

    @staticmethod
    def _parse_pt(value: str) -> float:
        """Converte '12pt' para float 12."""
        return float(re.sub(r'[^0-9.]', '', value) or '12')

    @staticmethod
    def _hex_to_rgb(hex_color: str) -> Tuple[int, int, int]:
        """Converte '#7030A0' para (112, 48, 160)."""
        hex_color = hex_color.lstrip('#')
        if len(hex_color) != 6:
            return (112, 48, 160)  # fallback Bravonix purple #7030A0
        return (
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16),
        )

    @staticmethod
    def _rgb_to_hex(rgb: Tuple[int, int, int]) -> str:
        """Converte (26, 54, 93) para '1a365d'."""
        return f'{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}'
