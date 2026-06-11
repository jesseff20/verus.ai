"""
DocumentProcessorService - Serviço para extração de texto de documentos.

Suporta extração de texto de PDFs e arquivos DOCX, com tratamento
de encoding e erros robustos.

SEGURANÇA: Valida magic bytes dos arquivos para prevenir uploads maliciosos.
"""
import logging
import re
import magic
from pathlib import Path
from typing import Dict, Optional, List
from io import BytesIO
from docx import Document
from pypdf import PdfReader
try:
    from odf.opendocument import load as odf_load
    from odf.text import P
    from odf import teletype
    HAS_ODFPY = True
except ImportError:
    HAS_ODFPY = False

logger = logging.getLogger(__name__)

# Magic bytes (assinaturas) dos tipos de arquivo permitidos
ALLOWED_MIME_TYPES = {
    'application/pdf': '.pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
    'text/plain': '.txt',
    'text/x-python': '.txt',  # Alguns .txt podem ser detectados assim
    'application/vnd.oasis.opendocument.text': '.odt',
    'application/octet-stream': None,  # Precisa validação adicional
}


class DocumentProcessorService:
    """
    Serviço para processar e extrair texto de documentos.

    Suporta:
    - PDF (usando pypdf)
    - DOCX (usando python-docx)
    - TXT (leitura direta)
    """

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.odt'}

    def extract_text(self, file_path: str) -> Dict[str, any]:
        """
        Extrai texto de um arquivo, detectando automaticamente o tipo.

        Args:
            file_path: Caminho para o arquivo

        Returns:
            Dict contendo:
                - text: Texto extraído
                - pages: Número de páginas (para PDF) ou None
                - file_type: Tipo do arquivo detectado
                - metadata: Metadados adicionais

        Raises:
            ValueError: Tipo de arquivo não suportado
            FileNotFoundError: Arquivo não encontrado
            Exception: Erro na extração
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")

        file_extension = path.suffix.lower()

        if file_extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Tipo de arquivo não suportado: {file_extension}. "
                f"Suportados: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        logger.info(f"Extraindo texto de {path.name} ({file_extension})")

        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension == '.docx':
                return self._extract_from_docx(file_path)
            elif file_extension == '.odt':
                return self._extract_from_odt(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
        except Exception as e:
            logger.error(f"Erro ao extrair texto de {path.name}: {str(e)}")
            raise

    def _extract_from_pdf(self, file_path: str) -> Dict[str, any]:
        """
        Extrai texto de um arquivo PDF.

        Args:
            file_path: Caminho para o arquivo PDF

        Returns:
            Dict com texto extraído e metadados
        """
        reader = PdfReader(file_path)
        num_pages = len(reader.pages)

        # Extrair texto de todas as páginas
        text_parts = []
        for page_num, page in enumerate(reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(f"--- Página {page_num} ---\n{page_text}")
            except Exception as e:
                logger.warning(f"Erro ao extrair texto da página {page_num}: {str(e)}")
                continue

        full_text = "\n\n".join(text_parts)

        # Extrair metadados
        metadata = {}
        if reader.metadata:
            metadata = {
                'title': reader.metadata.get('/Title', ''),
                'author': reader.metadata.get('/Author', ''),
                'subject': reader.metadata.get('/Subject', ''),
                'creator': reader.metadata.get('/Creator', ''),
            }

        result = {
            'text': full_text,
            'pages': num_pages,
            'file_type': 'pdf',
            'metadata': metadata,
            'char_count': len(full_text),
            'word_count': len(full_text.split())
        }

        logger.info(
            f"PDF extraído: {num_pages} páginas, "
            f"{result['char_count']} caracteres, "
            f"{result['word_count']} palavras"
        )

        return result

    def _extract_from_docx(self, file_path: str) -> Dict[str, any]:
        """
        Extrai texto de um arquivo DOCX.

        Args:
            file_path: Caminho para o arquivo DOCX

        Returns:
            Dict com texto extraído e metadados
        """
        doc = Document(file_path)

        # Extrair texto de todos os parágrafos
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)

        # Extrair texto de tabelas
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_texts.append(row_text)

        if table_texts:
            full_text += "\n\n--- TABELAS ---\n" + "\n".join(table_texts)

        # Metadados básicos
        metadata = {
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables)
        }

        # Tentar extrair propriedades do documento
        try:
            core_props = doc.core_properties
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
            })
        except Exception as e:
            logger.warning(f"Não foi possível extrair propriedades do DOCX: {str(e)}")

        result = {
            'text': full_text,
            'pages': None,  # DOCX não tem conceito de páginas
            'file_type': 'docx',
            'metadata': metadata,
            'char_count': len(full_text),
            'word_count': len(full_text.split())
        }

        logger.info(
            f"DOCX extraído: {metadata['paragraphs']} parágrafos, "
            f"{metadata['tables']} tabelas, "
            f"{result['char_count']} caracteres"
        )

        return result

    def _extract_from_odt(self, file_path: str) -> Dict[str, any]:
        """
        Extrai texto de um arquivo ODT usando odfpy.

        Args:
            file_path: Caminho para o arquivo ODT

        Returns:
            Dict com texto extraído e metadados
        """
        if not HAS_ODFPY:
            raise ImportError("odfpy não instalado. Execute: pip install odfpy")

        doc = odf_load(file_path)

        # Extrair todos os parágrafos
        paragraphs = doc.getElementsByType(P)
        text_parts = [teletype.extractText(p) for p in paragraphs]
        full_text = "\n\n".join(t for t in text_parts if t.strip())

        metadata = {
            'paragraphs': len(paragraphs),
        }

        result = {
            'text': full_text,
            'pages': None,
            'file_type': 'odt',
            'metadata': metadata,
            'char_count': len(full_text),
            'word_count': len(full_text.split())
        }

        logger.info(
            f"ODT extraído: {metadata['paragraphs']} parágrafos, "
            f"{result['char_count']} caracteres"
        )

        return result

    def _extract_from_txt(self, file_path: str) -> Dict[str, any]:
        """
        Extrai texto de um arquivo TXT.

        Args:
            file_path: Caminho para o arquivo TXT

        Returns:
            Dict com texto extraído e metadados
        """
        # Tentar diferentes encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()

                result = {
                    'text': text,
                    'pages': None,
                    'file_type': 'txt',
                    'metadata': {'encoding': encoding},
                    'char_count': len(text),
                    'word_count': len(text.split())
                }

                logger.info(
                    f"TXT extraído (encoding: {encoding}): "
                    f"{result['char_count']} caracteres"
                )

                return result

            except UnicodeDecodeError:
                continue

        raise ValueError(f"Não foi possível determinar o encoding do arquivo TXT")

    def validate_extraction(self, extraction_result: Dict[str, any]) -> Dict[str, any]:
        """
        Valida o resultado da extração de texto.

        Args:
            extraction_result: Resultado retornado por extract_text()

        Returns:
            Dict com:
                - is_valid: True se a extração é válida
                - issues: Lista de problemas encontrados
                - warnings: Lista de avisos
        """
        issues = []
        warnings = []

        text = extraction_result.get('text', '')

        # Verificar se há texto
        if not text or not text.strip():
            issues.append("Nenhum texto foi extraído do documento")

        # Verificar tamanho mínimo
        elif len(text) < 100:
            warnings.append(
                f"Texto muito curto ({len(text)} caracteres). "
                "Verifique se o documento foi extraído corretamente."
            )

        # Verificar se há caracteres estranhos (possível problema de encoding)
        weird_chars = sum(1 for c in text if ord(c) > 127 and c not in 'áéíóúàèìòùâêîôûãõçÁÉÍÓÚÀÈÌÒÙÂÊÎÔÛÃÕÇ')
        if weird_chars > len(text) * 0.1:  # Mais de 10% de caracteres estranhos
            warnings.append(
                "Documento contém muitos caracteres especiais. "
                "Pode haver problema de encoding."
            )

        return {
            'is_valid': len(issues) == 0,
            'issues': issues,
            'warnings': warnings
        }

    def chunk_text(
        self,
        text: str,
        chunk_size: int = 1000,
        overlap: int = 200
    ) -> list[str]:
        """
        Divide texto em chunks para processamento.

        Args:
            text: Texto a ser dividido
            chunk_size: Tamanho aproximado de cada chunk (em caracteres)
            overlap: Número de caracteres de sobreposição entre chunks

        Returns:
            Lista de chunks de texto
        """
        if not text:
            return []

        chunks = []
        start = 0

        while start < len(text):
            # Determinar fim do chunk
            end = start + chunk_size

            # Se não é o último chunk, tentar quebrar em ponto final ou nova linha
            if end < len(text):
                # Procurar ponto final no final do chunk
                last_period = text.rfind('.', start, end)
                last_newline = text.rfind('\n', start, end)

                break_point = max(last_period, last_newline)
                if break_point > start:
                    end = break_point + 1

            chunks.append(text[start:end].strip())
            start = max(end - overlap, start + 1)

        logger.info(f"Texto dividido em {len(chunks)} chunks")
        return chunks

    def _sanitize_filename(self, filename: str) -> str:
        """
        Sanitiza o nome do arquivo para prevenir path traversal e caracteres maliciosos.

        SEGURANÇA: Remove caracteres perigosos e previne ataques de path traversal.
        """
        # Remove path traversal
        filename = filename.replace('..', '').replace('/', '').replace('\\', '')

        # Remove caracteres especiais perigosos, mantém apenas alfanuméricos, -, _, .
        filename = re.sub(r'[^\w\-_\.]', '_', filename)

        # Remove múltiplos underscores consecutivos
        filename = re.sub(r'_+', '_', filename)

        # Limita tamanho do filename
        if len(filename) > 200:
            name, ext = filename.rsplit('.', 1) if '.' in filename else (filename, '')
            filename = name[:195] + ('.' + ext if ext else '')

        return filename

    def _validate_magic_bytes(self, file_content: bytes, claimed_extension: str) -> bool:
        """
        Valida o conteúdo real do arquivo usando magic bytes.

        SEGURANÇA: Previne upload de arquivos maliciosos disfarçados.

        Args:
            file_content: Primeiros bytes do arquivo
            claimed_extension: Extensão declarada pelo usuário

        Returns:
            True se o conteúdo corresponde à extensão, False caso contrário
        """
        try:
            # Detectar tipo real do arquivo
            detected_mime = magic.from_buffer(file_content, mime=True)
            logger.info(f"Magic bytes detectado: {detected_mime} para extensão {claimed_extension}")

            # Verificar se o MIME type é permitido
            if detected_mime not in ALLOWED_MIME_TYPES:
                logger.warning(f"SEGURANÇA: MIME type não permitido: {detected_mime}")
                return False

            # Verificar se a extensão corresponde ao conteúdo
            expected_extension = ALLOWED_MIME_TYPES.get(detected_mime)

            # application/octet-stream precisa validação adicional (pode ser DOCX ou ODT)
            if detected_mime == 'application/octet-stream':
                # DOCX e ODT são ZIPs, então podem ser detectados como octet-stream
                if claimed_extension in ('.docx', '.odt'):
                    # Verificar se começa com assinatura ZIP (PK)
                    if file_content[:2] == b'PK':
                        return True
                logger.warning(f"SEGURANÇA: octet-stream não validado para {claimed_extension}")
                return False

            # Para text/plain, aceita .txt
            if detected_mime in ('text/plain', 'text/x-python') and claimed_extension == '.txt':
                return True

            # Verificar correspondência
            if expected_extension and expected_extension != claimed_extension:
                logger.warning(
                    f"SEGURANÇA: Extensão não corresponde ao conteúdo! "
                    f"Declarado: {claimed_extension}, Detectado: {detected_mime}"
                )
                return False

            return True

        except Exception as e:
            logger.error(f"Erro ao validar magic bytes: {str(e)}")
            return False

    def process_uploaded_file(self, uploaded_file) -> Dict[str, any]:
        """
        Processa arquivo enviado via upload (Django UploadedFile).

        Args:
            uploaded_file: Instância de Django UploadedFile

        Returns:
            Dict com texto extraído, chunks e metadados

        Raises:
            ValueError: Se arquivo não é suportado ou inválido

        SEGURANÇA: Valida magic bytes e sanitiza filename.
        """
        # Sanitizar filename PRIMEIRO
        raw_filename = uploaded_file.name
        filename = self._sanitize_filename(raw_filename)

        if filename != raw_filename:
            logger.warning(f"SEGURANÇA: Filename sanitizado de '{raw_filename}' para '{filename}'")

        size = uploaded_file.size
        extension = Path(filename).suffix.lower()

        # Validar extensão
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Tipo de arquivo não suportado: {extension}. "
                f"Suportados: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        # Validar tamanho (50MB máximo)
        MAX_SIZE = 50 * 1024 * 1024
        if size > MAX_SIZE:
            raise ValueError(
                f"Arquivo muito grande: {size / 1024 / 1024:.1f}MB. "
                f"Máximo: {MAX_SIZE / 1024 / 1024}MB"
            )

        # SEGURANÇA: Validar magic bytes ANTES de processar
        file_content = uploaded_file.read()
        uploaded_file.seek(0)  # Reset para uso posterior

        if not self._validate_magic_bytes(file_content, extension):
            logger.error(f"SEGURANÇA: Upload bloqueado - magic bytes inválidos para {filename}")
            raise ValueError(
                f"Arquivo inválido: o conteúdo não corresponde ao tipo declarado ({extension}). "
                f"Possível tentativa de upload malicioso."
            )

        logger.info(f"Processando upload seguro: {filename} ({size} bytes)")

        # Processar baseado na extensão
        try:
            if extension == '.pdf':
                extraction = self._extract_from_pdf_upload(uploaded_file)
            elif extension == '.docx':
                extraction = self._extract_from_docx_upload(uploaded_file)
            elif extension == '.odt':
                extraction = self._extract_from_odt_upload(uploaded_file)
            elif extension == '.txt':
                extraction = self._extract_from_txt_upload(uploaded_file)
            else:
                raise ValueError(f"Extensão não reconhecida: {extension}")

            # Validar extração
            logger.info(f"[Debug] Passo 2: validando extração ({extraction['char_count']} chars)...")
            import sys; sys.stdout.flush()
            validation = self.validate_extraction(extraction)
            logger.info(f"[Debug] Passo 2: validação OK (valid={validation['is_valid']})")
            sys.stdout.flush()
            if not validation['is_valid']:
                raise ValueError(f"Extração inválida: {', '.join(validation['issues'])}")

            # Adicionar chunks
            logger.info(f"[Debug] Passo 3: chunking texto ({extraction['char_count']} chars)...")
            sys.stdout.flush()
            chunks = self.chunk_text(extraction['text'], chunk_size=1000, overlap=200)
            logger.info(f"[Debug] Passo 3: chunking OK ({len(chunks)} chunks)")
            sys.stdout.flush()

            logger.info(f"[Debug] Passo 4: montando resultado...")
            sys.stdout.flush()
            return {
                'filename': filename,
                'size': size,
                'file_type': extraction['file_type'],
                'text': extraction['text'],
                'chunks': chunks,
                'num_chunks': len(chunks),
                'char_count': extraction['char_count'],
                'word_count': extraction['word_count'],
                'metadata': extraction['metadata'],
                'warnings': validation['warnings']
            }

        except Exception as e:
            logger.error(f"Erro ao processar {filename}: {str(e)}")
            raise ValueError(f"Erro ao processar {filename}: {str(e)}")

    def _extract_from_pdf_upload(self, uploaded_file) -> Dict[str, any]:
        """Extrai texto de PDF enviado via upload."""
        logger.info(f"[PDF Debug] Início extração. type={type(uploaded_file).__name__}, size={getattr(uploaded_file, 'size', '?')}")

        uploaded_file.seek(0)
        raw = uploaded_file.read()
        logger.info(f"[PDF Debug] read() OK: {len(raw)} bytes lidos")

        reader = PdfReader(BytesIO(raw))
        num_pages = len(reader.pages)
        logger.info(f"[PDF Debug] PdfReader OK: {num_pages} páginas")

        # Liberar buffer bruto — pypdf já parseou
        del raw

        text_parts = []
        for page_num, page in enumerate(reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"[PDF Debug] Erro na página {page_num}: {str(e)}")
                continue

        full_text = "\n\n".join(text_parts)
        logger.info(f"[PDF Debug] Extração concluída: {len(full_text)} chars de {num_pages} páginas")

        metadata = {'pages': num_pages}
        if reader.metadata:
            metadata.update({
                'title': reader.metadata.get('/Title', ''),
                'author': reader.metadata.get('/Author', '')
            })

        return {
            'text': full_text,
            'pages': num_pages,
            'file_type': 'pdf',
            'metadata': metadata,
            'char_count': len(full_text),
            'word_count': len(full_text.split())
        }

    def _extract_from_docx_upload(self, uploaded_file) -> Dict[str, any]:
        """Extrai texto de DOCX enviado via upload."""
        doc = Document(BytesIO(uploaded_file.read()))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)

        # Extrair tabelas
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    full_text += "\n" + row_text

        metadata = {
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables)
        }

        return {
            'text': full_text,
            'pages': None,
            'file_type': 'docx',
            'metadata': metadata,
            'char_count': len(full_text),
            'word_count': len(full_text.split())
        }

    def _extract_from_odt_upload(self, uploaded_file) -> Dict[str, any]:
        """Extrai texto de ODT enviado via upload."""
        if not HAS_ODFPY:
            raise ImportError("odfpy não instalado. Execute: pip install odfpy")

        doc = odf_load(BytesIO(uploaded_file.read()))

        paragraphs = doc.getElementsByType(P)
        text_parts = [teletype.extractText(p) for p in paragraphs]
        full_text = "\n\n".join(t for t in text_parts if t.strip())

        metadata = {
            'paragraphs': len(paragraphs),
        }

        return {
            'text': full_text,
            'pages': None,
            'file_type': 'odt',
            'metadata': metadata,
            'char_count': len(full_text),
            'word_count': len(full_text.split())
        }

    def _extract_from_txt_upload(self, uploaded_file) -> Dict[str, any]:
        """Extrai texto de TXT enviado via upload."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        file_bytes = uploaded_file.read()

        for encoding in encodings:
            try:
                text = file_bytes.decode(encoding)
                return {
                    'text': text,
                    'pages': None,
                    'file_type': 'txt',
                    'metadata': {'encoding': encoding},
                    'char_count': len(text),
                    'word_count': len(text.split())
                }
            except UnicodeDecodeError:
                continue

        raise ValueError("Não foi possível decodificar o arquivo TXT")

    def process_multiple_uploads(self, uploaded_files: List) -> List[Dict]:
        """
        Processa múltiplos arquivos enviados via upload.

        Args:
            uploaded_files: Lista de Django UploadedFile

        Returns:
            Lista de dicts com informações de cada arquivo

        Raises:
            ValueError: Se algum arquivo falhar
        """
        results = []
        errors = []

        for file in uploaded_files:
            try:
                result = self.process_uploaded_file(file)
                results.append(result)
            except ValueError as e:
                errors.append(f"{file.name}: {str(e)}")
            except Exception as e:
                logger.error(f"Erro inesperado em {file.name}: {str(e)}")
                errors.append(f"{file.name}: Erro inesperado")

        if errors:
            raise ValueError(
                f"Erro ao processar {len(errors)} arquivo(s):\n" + "\n".join(errors)
            )

        logger.info(f"Processados {len(results)} arquivos com sucesso")
        return results
