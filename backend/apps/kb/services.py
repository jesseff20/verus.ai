"""
Serviços para processamento de documentos e busca vetorial
"""
import time
import logging
from typing import List, Dict, Optional
from django.db.models import F
from .models import Document, DocumentChunk
from apps.agents.services import EmbeddingService

logger = logging.getLogger(__name__)


class DocumentProcessingService:
    """Serviço para processamento de documentos"""

    @staticmethod
    def extract_text_from_bytes(file_content: bytes, file_type: str) -> str:
        """
        Extrai texto de conteúdo em memória (sem acessar storage/R2).

        Args:
            file_content: Bytes do arquivo
            file_type: Extensão do arquivo (pdf, docx, txt)

        Returns:
            Texto extraído
        """
        from io import BytesIO
        file_type = file_type.lower().replace('.', '')

        if file_type in ('txt', 'md', 'text/plain'):
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                return file_content.decode('latin-1')

        elif file_type in ('pdf', 'application/pdf'):
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(file_content))
            return '\n'.join(page.extract_text() for page in reader.pages)

        elif file_type in ('docx', 'doc', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword'):
            import docx
            doc = docx.Document(BytesIO(file_content))
            return '\n'.join(para.text for para in doc.paragraphs)

        elif file_type in ('odt', 'application/vnd.oasis.opendocument.text'):
            from odf.opendocument import load as odf_load
            from odf.text import P
            from odf import teletype
            doc = odf_load(BytesIO(file_content))
            paragraphs = doc.getElementsByType(P)
            return '\n'.join(teletype.extractText(p) for p in paragraphs if teletype.extractText(p).strip())

        else:
            raise ValueError(f"Tipo de arquivo não suportado: {file_type}")

    @staticmethod
    def extract_text(document: Document) -> str:
        """
        Extrai texto de documento (PDF, DOCX, TXT)

        Args:
            document: Instância do Document

        Returns:
            Texto extraído
        """
        file_type = document.file_type.lower()

        if file_type == 'text/plain' or file_type.endswith('txt'):
            return DocumentProcessingService._extract_text_txt(document)
        elif file_type == 'application/pdf' or file_type.endswith('pdf'):
            return DocumentProcessingService._extract_text_pdf(document)
        elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword'] or file_type.endswith('docx'):
            return DocumentProcessingService._extract_text_docx(document)
        elif file_type in ['application/vnd.oasis.opendocument.text'] or file_type.endswith('odt'):
            return DocumentProcessingService._extract_text_odt(document)
        else:
            raise ValueError(f"Tipo de arquivo não suportado: {file_type}")

    @staticmethod
    def _extract_text_txt(document: Document) -> str:
        """Extrai texto de arquivo TXT"""
        try:
            with document.file.open('r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            with document.file.open('r', encoding='latin-1') as f:
                return f.read()

    @staticmethod
    def _extract_text_pdf(document: Document) -> str:
        """Extrai texto de PDF"""
        try:
            from pypdf import PdfReader
            text = []
            with document.file.open('rb') as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text())
            return '\n'.join(text)
        except ImportError:
            raise ImportError("pypdf não instalado. Execute: pip install pypdf")
        except Exception as e:
            raise Exception(f"Erro ao extrair texto do PDF: {str(e)}")

    @staticmethod
    def _extract_text_docx(document: Document) -> str:
        """Extrai texto de DOCX"""
        try:
            import docx
            doc = docx.Document(document.file)
            return '\n'.join([para.text for para in doc.paragraphs])
        except ImportError:
            raise ImportError("python-docx não instalado. Execute: pip install python-docx")
        except Exception as e:
            raise Exception(f"Erro ao extrair texto do DOCX: {str(e)}")

    @staticmethod
    def _extract_text_odt(document: Document) -> str:
        """Extrai texto de ODT"""
        try:
            from odf.opendocument import load as odf_load
            from odf.text import P
            from odf import teletype
            doc = odf_load(document.file)
            paragraphs = doc.getElementsByType(P)
            return '\n'.join(teletype.extractText(p) for p in paragraphs if teletype.extractText(p).strip())
        except ImportError:
            raise ImportError("odfpy não instalado. Execute: pip install odfpy")
        except Exception as e:
            raise Exception(f"Erro ao extrair texto do ODT: {str(e)}")

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[Dict]:
        """
        Divide texto em chunks com overlap

        Args:
            text: Texto para dividir
            chunk_size: Tamanho aproximado de cada chunk (em caracteres)
            overlap: Overlap entre chunks (em caracteres)

        Returns:
            Lista de dicts com chunks e metadata
        """
        # Dividir por parágrafos primeiro
        paragraphs = text.split('\n')

        chunks = []
        current_chunk = ""
        chunk_index = 0

        for para in paragraphs:
            # Se adicionar este parágrafo ultrapassar o tamanho, criar novo chunk
            if len(current_chunk) + len(para) > chunk_size and current_chunk:
                chunks.append({
                    'content': current_chunk.strip(),
                    'chunk_index': chunk_index,
                    'metadata': {'length': len(current_chunk)}
                })

                # Criar overlap pegando últimos N caracteres
                current_chunk = current_chunk[-overlap:] + "\n" + para
                chunk_index += 1
            else:
                current_chunk += "\n" + para if current_chunk else para

        # Adicionar último chunk
        if current_chunk.strip():
            chunks.append({
                'content': current_chunk.strip(),
                'chunk_index': chunk_index,
                'metadata': {'length': len(current_chunk)}
            })

        return chunks

    @staticmethod
    def process_document(document: Document, file_content: bytes = None) -> int:
        """
        Processa documento completo: extrai texto, cria chunks e gera embeddings

        Args:
            document: Instância do Document
            file_content: Bytes do arquivo em memória (evita releitura do storage/R2)

        Returns:
            Número de chunks criados
        """
        timings = {}
        total_start = time.time()

        try:
            # 1. Extrair texto
            document.status = 'processing'
            document.save()

            t0 = time.time()
            if file_content:
                # Extração em memória - evita releitura do R2
                text = DocumentProcessingService.extract_text_from_bytes(
                    file_content, document.file_type
                )
                logger.info(f"=== Texto extraído da memória (sem R2) para {document.id} ===")
            else:
                text = DocumentProcessingService.extract_text(document)
            timings['1_extract_text'] = round(time.time() - t0, 3)

            # Validar se texto foi extraído
            if not text or not text.strip():
                raise Exception("Nenhum texto pôde ser extraído do documento. Verifique se o arquivo não está vazio ou corrompido.")

            t0 = time.time()
            document.extracted_text = text
            document.save()
            timings['2_save_text'] = round(time.time() - t0, 3)

            # 2. Criar chunks
            t0 = time.time()
            chunks_data = DocumentProcessingService.chunk_text(text, chunk_size=500, overlap=50)
            timings['3_chunk_text'] = round(time.time() - t0, 3)

            # Validar se chunks foram criados
            if not chunks_data or len(chunks_data) == 0:
                raise Exception("Não foi possível criar chunks do texto extraído. O documento pode ser muito curto.")

            # 3. Filtrar apenas chunks válidos (não vazios)
            valid_chunks = [chunk for chunk in chunks_data if chunk['content'].strip()]

            # Validar se há chunks válidos
            if not valid_chunks:
                raise Exception("Nenhum conteúdo válido encontrado nos chunks para gerar embeddings.")

            chunk_contents = [chunk['content'] for chunk in valid_chunks]

            # 4. Gerar embeddings (POSSÍVEL GARGALO)
            t0 = time.time()
            logger.info(f"[KB Pipeline] Gerando embeddings para {len(chunk_contents)} chunks...")
            embeddings = EmbeddingService.generate_batch(chunk_contents)
            timings['4_embeddings_api'] = round(time.time() - t0, 3)

            # 5. Criar DocumentChunks no banco (apenas para chunks válidos)
            t0 = time.time()
            chunk_objects = []
            for i, chunk_data in enumerate(valid_chunks):
                chunk_objects.append(
                    DocumentChunk(
                        document=document,
                        content=chunk_data['content'],
                        embedding=embeddings[i],
                        chunk_index=chunk_data['chunk_index'],
                        metadata=chunk_data['metadata']
                    )
                )
            timings['5_create_objects'] = round(time.time() - t0, 3)

            t0 = time.time()
            DocumentChunk.objects.bulk_create(chunk_objects)
            timings['6_bulk_create_db'] = round(time.time() - t0, 3)

            # 6. Marcar como processado
            t0 = time.time()
            document.status = 'completed'
            document.save()
            timings['7_final_save'] = round(time.time() - t0, 3)

            timings['TOTAL'] = round(time.time() - total_start, 3)

            # LOG ESTRUTURADO DE TIMING
            file_size_mb = round((document.file_size or 0) / 1024 / 1024, 2)
            logger.info(
                f"\n{'=' * 60}\n"
                f"[KB Pipeline] Documento processado: {document.title}\n"
                f"  Arquivo:       {file_size_mb}MB | Tipo: {document.file_type}\n"
                f"  Texto:         {len(text)} chars | {len(text.split())} palavras\n"
                f"  Chunks:        {len(chunk_contents)}\n"
                f"  ---\n"
                f"  Extração:      {timings['1_extract_text']}s\n"
                f"  Save texto DB: {timings['2_save_text']}s\n"
                f"  Chunking:      {timings['3_chunk_text']}s\n"
                f"  Embeddings:    {timings['4_embeddings_api']}s\n"
                f"  Criar objetos: {timings['5_create_objects']}s\n"
                f"  Bulk insert:   {timings['6_bulk_create_db']}s\n"
                f"  Save final:    {timings['7_final_save']}s\n"
                f"  ---\n"
                f"  TOTAL SYNC:    {timings['TOTAL']}s\n"
                f"{'=' * 60}"
            )

            return len(chunk_objects)

        except Exception as e:
            document.status = 'failed'
            document.processing_error = str(e)
            document.save()
            raise Exception(f"Erro ao processar documento: {str(e)}")


class VectorSearchService:
    """Serviço para busca vetorial usando pgvector"""

    @staticmethod
    def search(
        query_text: str,
        top_k: int = 5,
        similarity_threshold: float = 0.7,
        category: Optional[str] = None,
        tags: Optional[List[str]] = None,
        document_ids: Optional[List[str]] = None
    ) -> List[Dict]:
        """
        Busca semântica usando pgvector

        Args:
            query_text: Texto da query
            top_k: Quantidade de resultados
            similarity_threshold: Threshold mínimo de similaridade
            category: Filtrar por categoria
            tags: Filtrar por tags
            document_ids: Filtrar por IDs de documentos específicos

        Returns:
            Lista de chunks com scores de similaridade
        """
        # 1. Gerar embedding da query
        query_embedding = EmbeddingService.generate(query_text)

        # 2. Montar query com pgvector
        queryset = DocumentChunk.objects.select_related('document').filter(
            document__status__in=['completed', 'uploading', 'ready']
        )

        # Aplicar filtros
        if category:
            queryset = queryset.filter(document__category=category)

        if tags:
            # Filtrar documentos que contenham QUALQUER uma das tags
            from django.contrib.postgres.fields import ArrayField
            from django.db.models import Q
            tag_query = Q()
            for tag in tags:
                tag_query |= Q(document__tags__contains=[tag])
            queryset = queryset.filter(tag_query)

        if document_ids:
            queryset = queryset.filter(document_id__in=document_ids)

        # 3. Busca por similaridade usando pgvector
        # Usar cosine distance (operator <=>)
        from pgvector.django import CosineDistance

        queryset = queryset.annotate(
            distance=CosineDistance('embedding', query_embedding)
        ).filter(
            distance__lte=(1 - similarity_threshold)  # Converter threshold para distance
        ).order_by('distance')[:top_k]

        # 4. Formatar resultados
        results = []
        for chunk in queryset:
            # Converter distance para similarity score
            similarity = 1 - chunk.distance

            results.append({
                'chunk_id': str(chunk.id),
                'document_id': str(chunk.document_id),
                'document_title': chunk.document.title,
                'document_category': chunk.document.category,
                'content': chunk.content,
                'chunk_index': chunk.chunk_index,
                'similarity': round(similarity, 4),
                'metadata': chunk.metadata
            })

        return results
